"""ПР6 по Щеникову (ТСиУТИ) — Идентификация элементов математической модели.
Тема диссертации: разработка СППР по оценке рисков внедрения технологий ИИ
в IT-компании на основе нечёткой логики."""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


STUDENT_INITIALS = "Рослов Д.С."
GROUP = "М558М"
TEACHER = "Щеников Я.А."
DISCIPLINE = ("Теория систем и управление\n"
              "технологическими изменениями")
YEAR = "2026"
OUT = os.path.join(os.path.dirname(__file__), "ПР6_ТСиУТИ_Рослов.docx")

THESIS_TOPIC = ("Разработка системы поддержки принятия решений по оценке "
                "рисков внедрения технологий искусственного интеллекта в "
                "IT-компании на основе нечёткой логики")


def set_font(run, name="Times New Roman", size=14, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "start", "bottom", "end"):
        if edge in kwargs:
            attrs = kwargs[edge]
            tag = OxmlElement(f"w:{edge}")
            for k, v in attrs.items():
                tag.set(qn(f"w:{k}"), str(v))
            existing = tcBorders.find(qn(f"w:{edge}"))
            if existing is not None:
                tcBorders.remove(existing)
            tcBorders.append(tag)


def add_all_borders(table):
    border = {"sz": 4, "val": "single", "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)


def configure(doc):
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def add_par(doc, text, *, bold=False, italic=False, align=None, size=14,
            first_line_indent=None, line_spacing=1.5, space_after=0, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_body(doc, text):
    return add_par(
        doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25), line_spacing=1.5,
    )


def add_heading(doc, text):
    return add_par(
        doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
        size=14, line_spacing=1.5, space_before=6, space_after=6,
    )


def add_title_page(doc):
    configure(doc)
    add_par(doc, "САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ\nАЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ",
            bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, "Кафедра № 5\n«Инноватики и интегрированных систем качества»",
            align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)

    for _ in range(6):
        add_par(doc, "", size=12)

    add_par(doc, "ПРАКТИЧЕСКАЯ РАБОТА № 6",
            bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, "Идентификация элементов\nматематической модели",
            bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, f"по дисциплине «{DISCIPLINE}»",
            italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)

    for _ in range(6):
        add_par(doc, "", size=12)

    add_par(doc, f"Выполнил: студент гр. {GROUP}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)
    add_par(doc, STUDENT_INITIALS, align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, "Преподаватель:", align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)
    add_par(doc, TEACHER, align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)

    for _ in range(4):
        add_par(doc, "", size=12)

    add_par(doc, f"Санкт-Петербург\n{YEAR}",
            align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)

    doc.add_page_break()


def add_field(doc, label, body):
    """Заголовок-метка жирным + тело параграфа на той же строке."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    r = p.add_run(label + " ")
    set_font(r, size=14, bold=True)
    r = p.add_run(body)
    set_font(r, size=14)
    return p


def add_model_diagram(doc, caption):
    add_par(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12,
            space_before=6, space_after=3)
    headers = ["Вход", "Аппарат 1", "Аппарат 2", "Аппарат 3", "Выход"]
    bodies = [
        "Экспертные\nоценки рисков\n(матрица),\nпоказатели\nкомпании",
        "FMEA-анализ\n(идентификация\nи ранжирование\nрисков)",
        "Нечёткая система\nпринятия\nрешений\n(Mamdani)",
        "Дефаззификация\n(центр\nтяжести)",
        "Интегральный\nуровень риска\n(0–100) и\nрекомендация",
    ]
    t = doc.add_table(rows=2, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, size=11, bold=True)
    for j, b in enumerate(bodies):
        cell = t.rows[1].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        for i, line in enumerate(b.split("\n")):
            if i > 0:
                p.add_run().add_break()
            r = p.add_run(line)
            set_font(r, size=11)
    add_all_borders(t)


def main():
    doc = Document()
    add_title_page(doc)

    # Цель работы
    add_heading(doc, "ЦЕЛЬ РАБОТЫ")
    add_body(doc,
        "Получить навыки идентификации элементов математических моделей "
        "процессов и систем применительно к задаче собственного "
        "диссертационного исследования.")

    # Идентификация
    add_heading(doc, "ИДЕНТИФИКАЦИЯ ЭЛЕМЕНТОВ МАТЕМАТИЧЕСКОЙ МОДЕЛИ")

    add_field(doc, "Проблема (формулировка заказчика).",
        "Современные IT-компании активно внедряют технологии "
        "искусственного интеллекта в свои продукты и внутренние процессы. "
        "Однако решения о внедрении принимаются преимущественно на "
        "основе субъективных экспертных оценок без системного учёта "
        "технологических, организационных, кадровых, финансовых и "
        "регуляторных рисков. Это приводит к срыву сроков проектов, "
        "перерасходу бюджета, утечкам данных и репутационным потерям. "
        "Заказчик — руководство IT-компании — нуждается в инструменте, "
        "который позволил бы количественно оценивать совокупный риск "
        "внедрения той или иной ИИ-технологии и формировать "
        "обоснованную рекомендацию о целесообразности её внедрения.")

    add_field(doc, "Тема диссертационной работы.", THESIS_TOPIC + ".")

    add_field(doc, "Цель диссертационной работы.",
        "Повысить обоснованность управленческих решений о внедрении "
        "технологий искусственного интеллекта в IT-компании за счёт "
        "разработки системы поддержки принятия решений, обеспечивающей "
        "количественную оценку интегрального уровня риска. Цель "
        "измерима: ожидается повышение точности прогноза успешности "
        "внедрения ИИ-технологий с базового уровня (около 60 % "
        "по ретроспективным данным компании) до значения не менее 85 %, "
        "а также снижение доли проектов с превышением бюджета более "
        "чем на 20 % с 35 до 15 %.")

    add_field(doc, "Выход математической модели.",
        "Интегральный показатель риска R ∈ [0; 100] и соответствующая "
        "ему лингвистическая категория уровня риска (низкий, "
        "приемлемый, повышенный, высокий, критический), а также "
        "рекомендация СППР: «внедрять», «внедрять с компенсирующими "
        "мероприятиями», «отложить», «отказаться от внедрения». Выход "
        "представляет собой количественное выражение цели исследования.")

    add_field(doc, "Вход математической модели.",
        "Совокупность исходных данных по конкретному проекту внедрения "
        "ИИ-технологии: матрица экспертных оценок 9 идентифицированных "
        "рисков (технологических, организационных, кадровых, "
        "финансовых, регуляторных) по трём шкалам — вероятность "
        "возникновения, тяжесть последствий и вероятность обнаружения "
        "(значения от 1 до 10); статистические показатели компании "
        "(бюджет проекта, численность вовлечённого персонала, "
        "квалификация команды, наличие нормативной базы); качественные "
        "характеристики внедряемой технологии (тип ИИ-модели, степень "
        "зрелости, требования к данным). Часть входных данных носит "
        "числовой характер, часть — нечисловой (экспертные суждения, "
        "категории технологий).")

    add_field(doc, "Математический аппарат.",
        "Поскольку разработка модели с одним математическим аппаратом, "
        "удовлетворяющей всем требованиям задачи, не представляется "
        "возможной, в диссертационной работе предлагается гибридная "
        "математическая модель с тремя аппаратами, применяемыми "
        "последовательно: 1) модифицированный FMEA-анализ — "
        "идентифицирует риски, агрегирует экспертные оценки и "
        "формирует ранжированный вектор приоритетных чисел риска (RPN); "
        "2) нечёткая система принятия решений типа Мамдани — "
        "преобразует ранжированный вектор и качественные характеристики "
        "проекта в нечёткое множество уровней риска по базе из не менее "
        "чем 27 продукционных правил; 3) дефаззификация методом центра "
        "тяжести — переводит результат нечёткого вывода в чёткое "
        "значение интегрального показателя R ∈ [0; 100] и соответствующую "
        "категорию риска. Выбор аппаратов соответствует таблице 6.2 "
        "методических указаний: для входа смешанного типа (числовой "
        "вектор + нечисловые данные) и числового выхода обоснованным "
        "является применение нечёткой системы принятия решений в "
        "сочетании с экспертными методами.")

    add_model_diagram(doc,
        "Рисунок 1 — Структура гибридной математической модели СППР")

    add_field(doc, "Верификация математической модели.",
        "Верификация выполняется на ретроспективных данных IT-компании. "
        "Из архива выбираются не менее 30 завершённых проектов "
        "внедрения ИИ-технологий за последние 3–5 лет с известными "
        "фактическими результатами (успешно / частично успешно / "
        "неуспешно, по совокупности критериев — соблюдение сроков, "
        "бюджета, целевых показателей качества). По каждому проекту "
        "восстанавливаются входные данные модели, вычисляется "
        "прогнозный интегральный показатель риска R, после чего "
        "результат сопоставляется с фактическим исходом. Для оценки "
        "качества прогноза используются метрики точности классификации "
        "(accuracy, precision, recall) и коэффициент корреляции Пирсона "
        "между прогнозным R и фактическим показателем затрат проекта. "
        "Дополнительно проводится валидация модели методом "
        "анкетирования группы из не менее 5 экспертов компании: "
        "эксперты оценивают полученные моделью рекомендации по "
        "пятибалльной шкале, итоговая средняя оценка должна составлять "
        "не менее 4,0 баллов. Имитационный эксперимент строится на "
        "сценариях типа «что-если» — исследуется реакция модели на "
        "вариации экспертных оценок в пределах ±20 % для проверки "
        "устойчивости выхода.")

    # Заключение
    add_heading(doc, "ВЫВОДЫ")
    add_body(doc,
        "В работе выполнена идентификация элементов математической "
        "модели для собственного диссертационного исследования: "
        "сформулированы проблема, тема и цель, определены вход и выход "
        "будущей модели, выбран и обоснован математический аппарат — "
        "гибридная модель из трёх последовательно применяемых "
        "аппаратов (FMEA-анализ, нечёткая система принятия решений "
        "Мамдани, дефаззификация). Определён план верификации модели, "
        "включающий ретроспективный эксперимент на архивных данных "
        "IT-компании, оценку точности классификации и валидацию "
        "методом экспертного анкетирования. Полученные результаты "
        "соответствуют требованиям закона полноты частей системы и "
        "закона энергетической проводимости: модель содержит все три "
        "обязательных элемента (вход, аппарат, выход), а сквозной "
        "поток данных от входа к выходу обеспечен преемственностью "
        "форматов между аппаратами.")

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
