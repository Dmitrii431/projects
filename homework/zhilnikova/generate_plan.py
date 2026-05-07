"""ПР2 по Жильниковой — Составление плана ВКР.
Тема диссертации: разработка СППР по оценке рисков внедрения технологий ИИ
в IT-компании на основе нечёткой логики."""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


STUDENT_INITIALS = "Рослов Д.С."
GROUP = "М558М"
TEACHER = "Жильникова Н.А."
DISCIPLINE = "Научно-исследовательская работа"
YEAR = "2026"
OUT = os.path.join(os.path.dirname(__file__), "План_ВКР_Рослов.docx")

THESIS_TOPIC = ("Разработка системы поддержки принятия решений по оценке "
                "рисков внедрения технологий искусственного интеллекта в "
                "IT-компании на основе нечёткой логики")


def set_font(run, name="Times New Roman", size=14, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def configure(doc):
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def add_par(doc, text, *, bold=False, italic=False, align=None, size=14,
            first_line_indent=None, left_indent=None, line_spacing=1.5,
            space_after=0, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_centered(doc, text, *, bold=False, italic=False, size=14, line_spacing=1.5,
                 space_before=0, space_after=0):
    return add_par(doc, text, bold=bold, italic=italic,
                   align=WD_ALIGN_PARAGRAPH.CENTER, size=size,
                   line_spacing=line_spacing, space_before=space_before,
                   space_after=space_after)


def add_title_page(doc):
    configure(doc)
    add_centered(doc,
        "САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ\n"
        "АЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ",
        bold=True, size=14, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_centered(doc, "Кафедра № 5\n«Инноватики и интегрированных систем качества»",
                 size=14, line_spacing=1.15)

    for _ in range(6):
        add_par(doc, "", size=12)

    add_centered(doc, "ОЦЕНКА ПРАКТИЧЕСКОЙ РАБОТЫ",
                 bold=True, size=16, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_centered(doc,
        "Составление плана\nвыпускной квалификационной работы",
        bold=True, size=16, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_centered(doc, f"по дисциплине «{DISCIPLINE}»",
                 italic=True, size=14, line_spacing=1.15)

    for _ in range(7):
        add_par(doc, "", size=12)

    add_par(doc, f"Выполнил: студент гр. {GROUP}",
            align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)
    add_par(doc, STUDENT_INITIALS,
            align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, "Преподаватель:",
            align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)
    add_par(doc, f"{TEACHER}, профессор, д.т.н., доц.",
            align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)

    for _ in range(4):
        add_par(doc, "", size=12)

    add_centered(doc, f"Санкт-Петербург\n{YEAR}",
                 size=14, line_spacing=1.15)
    doc.add_page_break()


def add_section_heading(doc, text):
    return add_par(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                   size=14, line_spacing=1.5, space_before=12, space_after=6)


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    r = p.add_run(label + " ")
    set_font(r, size=14, bold=True)
    r = p.add_run(value)
    set_font(r, size=14)
    return p


def add_chapter_title(doc, text):
    return add_par(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   size=14, line_spacing=1.5, space_before=10, space_after=4)


def add_paragraph_item(doc, text):
    return add_par(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   size=14, line_spacing=1.5,
                   left_indent=Cm(1.25), space_after=2)


def add_simple(doc, text):
    return add_par(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=14,
                   line_spacing=1.5, space_before=4, space_after=2)


def main():
    doc = Document()
    add_title_page(doc)

    # 1. Тема ВКР
    add_section_heading(doc, "1. Тема выпускной квалификационной работы")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    r = p.add_run("«" + THESIS_TOPIC + "».")
    set_font(r, size=14, italic=True)

    add_par(doc,
        "Тема сформулирована таким образом, чтобы максимально конкретно "
        "отразить основную идею работы: объект (риски внедрения "
        "технологий искусственного интеллекта в IT-компании), главный "
        "результат (система поддержки принятия решений) и метод (аппарат "
        "нечёткой логики).",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=14, line_spacing=1.5,
        first_line_indent=Cm(1.25), space_before=4)

    # 2. Научный руководитель
    add_section_heading(doc, "2. Научный руководитель")
    add_label_value(doc, "Ф. И. О.:", "Гетманова Галина Владимировна.")
    add_label_value(doc, "Должность:", "доцент кафедры № 5 «Инноватики и "
                    "интегрированных систем качества» Института "
                    "фундаментальной подготовки и технологических инноваций "
                    "(Институт ФПТИ).")
    add_label_value(doc, "Учёная степень:", "кандидат экономических наук.")
    add_label_value(doc, "Учёное звание:", "доцент.")
    add_label_value(doc, "Контакты:", "БМ, ауд. 23-23; getmanova@guap.ru.")
    add_label_value(doc, "Научные интересы:",
        "управленческие инновации и инновационный менеджмент; "
        "ситуационное управление качеством и инновационной деятельностью "
        "в сложных технических системах; стратегический менеджмент; "
        "методы управления проектами; организационные структуры "
        "управления; диагностика и внедрение управленческих инноваций в "
        "деятельности предприятия (тема кандидатской диссертации, 2005).")

    # 3. План ВКР
    add_section_heading(doc, "3. Развёрнутый план ВКР")

    add_simple(doc, "Введение")

    add_chapter_title(doc,
        "Глава 1. Теоретические основы оценки рисков внедрения технологий "
        "искусственного интеллекта в IT-компаниях")
    for it in (
        "1.1. Современное состояние и тенденции развития технологий "
        "искусственного интеллекта в IT-секторе.",
        "1.2. Классификация и характеристика рисков внедрения "
        "ИИ-технологий: технологические, организационные, финансовые, "
        "регуляторные и этические.",
        "1.3. Обзор существующих методов и моделей оценки рисков внедрения "
        "новых технологий.",
        "1.4. Анализ применимости аппарата нечёткой логики к задачам "
        "оценки рисков в условиях неопределённости.",
        "1.5. Выводы по главе 1.",
    ):
        add_paragraph_item(doc, it)

    add_chapter_title(doc,
        "Глава 2. Анализ практики внедрения технологий искусственного "
        "интеллекта и оценки сопутствующих рисков (на примере IT-компаний)")
    for it in (
        "2.1. Организационная характеристика и анализ практики внедрения "
        "ИИ-технологий в исследуемых IT-компаниях.",
        "2.2. Идентификация и систематизация рисков ИИ-проектов на основе "
        "ретроспективного анализа архивных данных.",
        "2.3. Сравнительный анализ применяемых методик оценки рисков и "
        "оценка их адекватности задачам IT-сектора.",
        "2.4. Выявление пробелов в существующих подходах и обоснование "
        "требований к гибридной модели оценки рисков.",
        "2.5. Выводы по главе 2.",
    ):
        add_paragraph_item(doc, it)

    add_chapter_title(doc,
        "Глава 3. Разработка системы поддержки принятия решений по оценке "
        "рисков внедрения ИИ-технологий на основе нечёткой логики")
    for it in (
        "3.1. Разработка архитектуры гибридной СППР: модифицированный "
        "FMEA-анализ, нечёткий вывод Мамдани, дефаззификация методом "
        "центра тяжести.",
        "3.2. Формирование лингвистических переменных, базы продукционных "
        "правил и калибровка функций принадлежности.",
        "3.3. Программная реализация СППР и верификация модели на "
        "ретроспективных данных IT-компаний.",
        "3.4. Технико-экономическое обоснование внедрения СППР и оценка "
        "социально-экономического эффекта.",
        "3.5. Выводы по главе 3.",
    ):
        add_paragraph_item(doc, it)

    add_simple(doc, "Заключение")
    add_simple(doc, "Список использованных источников")
    add_simple(doc, "Приложения")

    # 4. Календарный план
    add_section_heading(doc, "4. Ориентировочный календарный план выполнения ВКР")
    add_par(doc,
        "Защита ВКР запланирована на июнь 2027 г. Распределение этапов "
        "работы по семестрам приведено ниже.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=14, line_spacing=1.5,
        first_line_indent=Cm(1.25), space_after=4)
    for it in (
        "Май–август 2026 г. — формирование темы, поиск и анализ "
        "литературы, патентный поиск, написание научной статьи "
        "по результатам обзора (научно-исследовательская работа).",
        "Сентябрь–декабрь 2026 г. — подготовка теоретической части "
        "(глава 1), уточнение методологии исследования, выбор "
        "математического аппарата.",
        "Январь–апрель 2027 г. — сбор и анализ эмпирических данных "
        "(глава 2), идентификация рисков, проектирование архитектуры "
        "СППР, разработка лингвистических переменных и базы правил.",
        "Апрель–май 2027 г. — программная реализация СППР, верификация "
        "и валидация модели на ретроспективных данных, технико-"
        "экономическое обоснование (глава 3), оформление заключения и "
        "приложений.",
        "Май 2027 г. — оформление ВКР по требованиям ГОСТ, проверка на "
        "оригинальность, прохождение нормоконтроля, подготовка "
        "презентации и доклада.",
        "Июнь 2027 г. — предзащита, доработка по замечаниям, защита ВКР.",
    ):
        add_paragraph_item(doc, it)

    # 5. Обоснование структуры
    add_section_heading(doc, "5. Обоснование логики плана")
    add_par(doc,
        "Структура плана соответствует классической логике "
        "диссертационного исследования и обеспечивает поэтапное "
        "раскрытие темы. Первая глава носит теоретический характер и "
        "формирует методологическую базу: рассматриваются текущее "
        "состояние ИИ-технологий, их классификация, существующие методы "
        "оценки рисков и возможности применения нечёткой логики. Вторая "
        "глава является аналитической и опирается на эмпирический "
        "материал — практику внедрения ИИ-технологий в IT-компаниях; в "
        "результате выявляются пробелы в существующих подходах и "
        "формулируются требования к разрабатываемой модели. Третья "
        "глава — практическая и проектная: на основе результатов первых "
        "двух глав разрабатывается архитектура гибридной СППР, "
        "выполняется её программная реализация и верификация, проводится "
        "технико-экономическое обоснование внедрения. Главы и параграфы "
        "сбалансированы по объёму, что соответствует требованиям к "
        "оформлению ВКР. Объекты, предметы и методы исследования, "
        "формулируемые во введении, последовательно прорабатываются на "
        "всех уровнях плана, обеспечивая его логическую целостность.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=14, line_spacing=1.5,
        first_line_indent=Cm(1.25))

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
