"""ПР3 по Жильниковой — Разработка введения к магистерской диссертации.
Тема: разработка СППР по оценке рисков внедрения технологий ИИ
в IT-компании на основе нечёткой логики."""
import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


STUDENT_INITIALS = "Рослов Д.С."
GROUP = "М558М"
TEACHER = "Жильникова Н.А."
DISCIPLINE = "Научно-исследовательская работа"
YEAR = "2026"
OUT = os.path.join(os.path.dirname(__file__), "Введение_ВКР_Рослов.docx")

THESIS_TOPIC = ("Разработка системы поддержки принятия решений по оценке "
                "рисков внедрения технологий искусственного интеллекта в "
                "IT-компании на основе нечёткой логики")


def set_font(run, name="Times New Roman", size=14, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def configure(doc):
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def add_par(doc, text, *, bold=False, italic=False, align=None, size=14,
            first_line_indent=None, left_indent=None, line_spacing=1.5,
            space_after=0, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_centered(doc, text, *, bold=False, italic=False, size=14, line_spacing=1.5,
                 space_before=0, space_after=0):
    return add_par(doc, text, bold=bold, italic=italic,
                   align=WD_ALIGN_PARAGRAPH.CENTER, size=size,
                   line_spacing=line_spacing, space_before=space_before,
                   space_after=space_after)


def _set_cell_borders(cell, top=False, bottom=False, no_others=True):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        if (edge == "top" and top) or (edge == "bottom" and bottom):
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")


def _cell_par(cell, text, *, align=WD_ALIGN_PARAGRAPH.CENTER, size=12,
              italic=False, line_spacing=1.0, space_after=0):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    r = p.add_run(text)
    set_font(r, size=size, italic=italic)
    return p


def add_signature_block(doc, *, top_left, top_right, bottom_left, bottom_middle,
                        bottom_right, widths=(6.0, 5.5, 5.0)):
    """Three-column block: top row with values (above the line), bottom row
    with line + caption underneath. Middle column = signature/date (empty)."""
    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)
    # Top row (values above the line)
    top = [top_left, "", top_right]
    for i, val in enumerate(top):
        cell = table.rows[0].cells[i]
        _set_cell_borders(cell)
        _cell_par(cell, val, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    # Bottom row (caption beneath the line)
    bot = [bottom_left, bottom_middle, bottom_right]
    for i, val in enumerate(bot):
        cell = table.rows[1].cells[i]
        _set_cell_borders(cell, top=True)
        _cell_par(cell, val, align=WD_ALIGN_PARAGRAPH.CENTER, size=10,
                  italic=True)
    return table


def add_title_page(doc):
    configure(doc)
    add_centered(doc,
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ "
        "ФЕДЕРАЦИИ",
        size=14, line_spacing=1.15)
    add_centered(doc,
        "федеральное государственное автономное образовательное "
        "учреждение высшего образования",
        size=12, line_spacing=1.15)
    add_centered(doc,
        "«САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ "
        "АЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ»",
        size=14, line_spacing=1.15)

    for _ in range(2):
        add_par(doc, "", size=12)

    add_centered(doc, "КАФЕДРА №5", size=14, line_spacing=1.15)

    for _ in range(3):
        add_par(doc, "", size=12)

    add_par(doc, "ОЦЕНКА ПРАКТИЧЕСКОЙ РАБОТЫ",
            align=WD_ALIGN_PARAGRAPH.LEFT, size=14, line_spacing=1.15)
    add_par(doc, "ПРЕПОДАВАТЕЛЬ",
            align=WD_ALIGN_PARAGRAPH.LEFT, size=14, line_spacing=1.15)

    add_signature_block(doc,
        top_left="Профессор, д.т.н., доц.",
        top_right="Н. А. Жильникова",
        bottom_left="должность, уч. степень,\nзвание",
        bottom_middle="подпись, дата",
        bottom_right="инициалы, фамилия")

    for _ in range(2):
        add_par(doc, "", size=12)

    add_centered(doc,
        "Разработка введения к магистерской диссертации",
        size=14, line_spacing=1.15)

    add_par(doc, "", size=12)

    add_centered(doc, f"по дисциплине: {DISCIPLINE}",
                 size=14, line_spacing=1.15)

    for _ in range(3):
        add_par(doc, "", size=12)

    add_par(doc, "РАБОТУ ВЫПОЛНИЛ(А)",
            align=WD_ALIGN_PARAGRAPH.LEFT, size=14, line_spacing=1.15)
    add_par(doc, f"СТУДЕНТ(КА) ГР. № {GROUP}",
            align=WD_ALIGN_PARAGRAPH.LEFT, size=14, line_spacing=1.15)

    add_signature_block(doc,
        top_left="",
        top_right=STUDENT_INITIALS,
        bottom_left="",
        bottom_middle="подпись, дата",
        bottom_right="инициалы, фамилия")

    for _ in range(3):
        add_par(doc, "", size=12)

    add_centered(doc, f"Санкт-Петербург {YEAR}",
                 size=14, line_spacing=1.15)
    doc.add_page_break()


def add_section_heading(doc, text):
    return add_par(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   size=14, line_spacing=1.5, space_before=10, space_after=4,
                   first_line_indent=Cm(1.25))


def add_body(doc, text):
    return add_par(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=14,
                   line_spacing=1.5, first_line_indent=Cm(1.25), space_after=2)


def add_list_item(doc, text):
    return add_par(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=14,
                   line_spacing=1.5, first_line_indent=Cm(1.25), space_after=2)


def main():
    doc = Document()
    add_title_page(doc)

    # Заголовок «ВВЕДЕНИЕ»
    add_centered(doc, "ВВЕДЕНИЕ", bold=True, size=14, line_spacing=1.5,
                 space_after=8)

    # ========== 1. АКТУАЛЬНОСТЬ ==========
    add_section_heading(doc, "Актуальность исследования.")
    add_body(doc,
        "Технологии искусственного интеллекта (ИИ) стали ключевым "
        "фактором цифровой трансформации мировой экономики. По данным "
        "McKinsey Global Institute и Stanford AI Index 2024, ежегодные "
        "инвестиции в ИИ-решения превышают 250 млрд долл. США, при этом "
        "доля IT-сектора в общем объёме внедрений превышает 35 %. В "
        "Российской Федерации развитие ИИ закреплено как национальный "
        "приоритет: Указ Президента РФ от 10.10.2019 № 490 утвердил "
        "Национальную стратегию развития ИИ до 2030 года, объём "
        "российского рынка ИИ-технологий в 2024 г. превысил 900 млрд "
        "руб. с прогнозом роста до 1,7 трлн руб. к 2026 г.")
    add_body(doc,
        "Однако фактическая результативность ИИ-проектов остаётся "
        "низкой: по оценкам Gartner и IDC, от 70 до 85 % инициатив не "
        "достигают заявленных целей. Причина — сочетание специфических "
        "рисков, нехарактерных для классических IT-проектов: "
        "технологическая неопределённость, смещённость обучающих "
        "данных, нестабильность моделей (галлюцинации, concept drift), "
        "регуляторные и этические ограничения, угрозы информационной "
        "безопасности (adversarial-атаки, утечка персональных данных), "
        "организационное сопротивление и кадровый дефицит.")
    add_body(doc,
        "Существующие методы (FMEA, экспертные методы, вероятностный "
        "анализ) ориентированы на количественные данные и плохо "
        "работают при малых выборках прецедентов и качественном "
        "характере большинства факторов. Нечёткая логика Л. А. Заде "
        "является признанным аппаратом обработки качественной "
        "информации в условиях неопределённости, однако существующие "
        "нечёткие модели редко учитывают отраслевую специфику ИИ-"
        "проектов в IT-секторе.")
    add_body(doc,
        "Таким образом, актуальность исследования обусловлена "
        "стратегической значимостью внедрения ИИ-технологий, высокой "
        "долей неуспешных ИИ-проектов и недостаточной разработанностью "
        "методического и инструментального обеспечения для оценки "
        "связанных с ними рисков в IT-компаниях.")

    # ========== 2. СОСТОЯНИЕ НАУЧНОЙ РАЗРАБОТАННОСТИ ==========
    add_section_heading(doc, "Состояние научной разработанности проблемы.")
    add_body(doc,
        "Теоретические основы нечёткой логики заложены в трудах "
        "Л. А. Заде, развиты Е. Мамдани, М. Сугено, Т. Такаги. "
        "Российская школа представлена работами Д. А. Поспелова, "
        "А. Н. Борисова, А. В. Леоненкова, Н. Г. Ярушкиной, "
        "А. П. Ротштейна по нечёткому моделированию и СППР. "
        "Проблематика управления рисками раскрыта в работах "
        "К. З. Билятдинова, Г. В. Гетмановой, стандартах ISO "
        "31000:2018, COSO ERM, PMBoK Guide; применение FMEA — в "
        "трудах Дж. Боулза, Х. Лю, J. Mendel.")
    add_body(doc,
        "Вопросы управления рисками ИИ-систем рассматриваются в "
        "международных документах NIST AI RMF 1.0 (2023), ISO/IEC "
        "23894:2023, ISO/IEC 42001:2023, а также в работах "
        "Дж. Мёкандера, Р. Бинса и групп IBM Responsible AI, "
        "Microsoft, MIT-IBM Watson Lab. Отдельные приложения "
        "нечёткой логики к оценке IT-рисков представлены у "
        "Е. С. Назаревича, И. А. Вершининой, Г. В. Гетмановой.")
    add_body(doc,
        "Выполненный автором систематический обзор по протоколу "
        "PRISMA (134 публикации и 47 патентов за 2015–2025 гг.) "
        "показывает, что комплексные СППР, объединяющие "
        "модифицированный FMEA-анализ с нечётким выводом Мамдани "
        "и откалиброванные на отраслевых данных IT-компаний, в "
        "открытой литературе практически не представлены. Именно "
        "это «белое пятно» — разработка такой гибридной СППР с "
        "учётом отраслевой специфики ИИ-проектов в IT — "
        "рассматривается в настоящем исследовании.")

    # ========== 3. ОБЪЕКТ И ПРЕДМЕТ ==========
    add_section_heading(doc, "Объект и предмет исследования.")
    add_body(doc,
        "Объектом исследования являются процессы внедрения технологий "
        "искусственного интеллекта в IT-компаниях.")
    add_body(doc,
        "Предмет исследования — методы и средства оценки рисков "
        "внедрения технологий искусственного интеллекта на основе "
        "аппарата нечёткой логики.")

    # ========== 4. ЦЕЛЬ И ЗАДАЧИ ==========
    add_section_heading(doc, "Цель и задачи исследования.")
    add_body(doc,
        "Цель исследования — разработать систему поддержки принятия "
        "решений по оценке рисков внедрения технологий искусственного "
        "интеллекта в IT-компании на основе аппарата нечёткой логики.")
    add_body(doc,
        "Для достижения цели сформулированы задачи:")
    for it in (
        "1. Проанализировать современное состояние ИИ-технологий в "
        "IT-секторе, идентифицировать и классифицировать характерные "
        "риски их внедрения.",
        "2. Выполнить систематический обзор существующих методов "
        "оценки рисков и обосновать применимость аппарата нечёткой "
        "логики.",
        "3. Разработать гибридную модель оценки рисков, включающую "
        "модифицированный FMEA-анализ, нечёткий вывод Мамдани и "
        "дефаззификацию методом центра тяжести.",
        "4. Сформировать систему лингвистических переменных и базу "
        "продукционных правил, откалиброванную на ретроспективных "
        "данных IT-проектов.",
        "5. Реализовать СППР программно, провести её верификацию и "
        "валидацию, выполнить технико-экономическое обоснование "
        "внедрения.",
    ):
        add_list_item(doc, it)

    # ========== 5. ТЕОРЕТИЧЕСКАЯ И МЕТОДОЛОГИЧЕСКАЯ БАЗА ==========
    add_section_heading(doc,
        "Теоретическая и методологическая база исследования.")
    add_body(doc,
        "Теоретическую базу составляют: теория нечётких множеств и "
        "нечёткого вывода (Л. А. Заде, Е. Мамдани, М. Сугено); "
        "теория управления рисками (К. З. Билятдинов, ISO 31000, "
        "COSO ERM); методология FMEA (IEC 60812); концепции "
        "ответственного ИИ (NIST AI RMF, ISO/IEC 23894, 42001); "
        "работы по проектированию СППР (А. Н. Борисов, "
        "А. В. Леоненков, Н. Г. Ярушкина).")
    add_body(doc,
        "Методологическая база включает: систематический обзор по "
        "протоколу PRISMA и патентный поиск (Роспатент, Espacenet, "
        "USPTO, WIPO) — для оценки разработанности проблемы; "
        "модифицированный FMEA-анализ — для идентификации и "
        "приоритизации рисков; нечёткое моделирование и вывод "
        "Мамдани — для интегральной оценки рисков в условиях "
        "неопределённости; имитационное моделирование — для "
        "верификации; экспертные оценки с расчётом коэффициента "
        "конкордации Кендалла — для калибровки функций "
        "принадлежности; ретроспективный сравнительный анализ — для "
        "валидации СППР на исторических данных IT-проектов.")

    # ========== 6. ПРАКТИЧЕСКАЯ ЗНАЧИМОСТЬ ==========
    add_section_heading(doc, "Практическая значимость исследования.")
    add_body(doc,
        "Практическая значимость состоит в следующем: разработанная "
        "СППР позволяет проводить интегральную оценку рисков ИИ-"
        "проектов на ранней стадии и снижать вероятность их "
        "неуспеха; сформированная база продукционных правил может "
        "использоваться как методическая основа внутреннего риск-"
        "менеджмента и аудита ИИ-инициатив в IT-компаниях; "
        "результаты применимы при разработке отраслевых стандартов "
        "и рекомендаций для регуляторов (Минцифры России, АНО "
        "«Цифровая экономика»); материалы исследования внедряются "
        "в учебный процесс ГУАП по направлению 27.04.05 «Инноватика».")

    # ========== 7. НОВИЗНА ==========
    add_section_heading(doc, "Научная новизна исследования.")
    add_body(doc,
        "Научная новизна заключается в следующем:")
    for it in (
        "1. Предложена авторская классификация рисков внедрения ИИ-"
        "технологий в IT-компаниях, учитывающая специфические "
        "факторы (галлюцинации, concept drift, adversarial-атаки, "
        "регуляторно-этические ограничения).",
        "2. Разработана гибридная модель, объединяющая "
        "модифицированный FMEA с нечётким выводом Мамдани и "
        "дефаззификацией методом центра тяжести; получаемый "
        "интегральный риск R∈[0;100] сопровождается категорией и "
        "рекомендацией СППР.",
        "3. Сформирована откалиброванная база продукционных правил "
        "и предложена методика валидации СППР на ретроспективных "
        "данных с применением метрик accuracy, precision, recall и "
        "экспертной шкалы согласованности.",
    ):
        add_list_item(doc, it)

    # ========== 8. АПРОБАЦИЯ ==========
    add_section_heading(doc, "Апробация результатов исследования.")
    add_body(doc,
        "Основные положения диссертации докладывались на научно-"
        "исследовательских семинарах кафедры № 5 «Инноватики и "
        "интегрированных систем качества» ГУАП (Санкт-Петербург, "
        "2026 г.). Результаты планируются к представлению на "
        "ежегодной научно-технической конференции ГУАП и "
        "международной научно-практической конференции "
        "«Инновационная экономика и менеджмент: методы и технологии» "
        "(Санкт-Петербург, 2027 г.). По теме диссертации "
        "подготовлена научная публикация:")
    add_list_item(doc,
        "Рослов Д. С. Современное состояние и перспективы применения "
        "нечёткой логики для оценки рисков внедрения технологий "
        "искусственного интеллекта в IT-компаниях: систематический "
        "аналитический обзор / Д. С. Рослов ; науч. рук. "
        "Г. В. Гетманова. — Санкт-Петербург : ГУАП, 2026. — 16 с.")
    add_body(doc,
        "Дальнейшая апробация планируется в виде публикаций в "
        "изданиях из перечня ВАК и базы РИНЦ.")

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
