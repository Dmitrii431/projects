#!/usr/bin/env python3
"""Сборка Word-документа с ответами к экзамену по ТСиУТИ (Щеников Я.А.)
из markdown-файлов глав. Титульный лист ГУАП, кафедра №5."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

BASE = Path("/home/sergei/dima-projects/homework/shenikov_exam")
ANSW = BASE / "answers"
OUT = BASE / "Ответы_к_экзамену_ТСиУТИ_Рослов.docx"

FONT = "Times New Roman"

doc = Document()

# базовый стиль
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(14)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(0)

# поля
sec = doc.sections[0]
sec.left_margin = Cm(3)
sec.right_margin = Cm(1.5)
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)


def set_font(run, size=14, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)


def add_runs(par, text, size=14, base_bold=False):
    """Разбор inline **жирного** на runs."""
    for i, chunk in enumerate(re.split(r"\*\*", text)):
        if chunk == "":
            continue
        run = par.add_run(chunk)
        set_font(run, size=size, bold=base_bold or (i % 2 == 1))


# ---------- ТИТУЛЬНЫЙ ЛИСТ ----------
def center(text, size=14, bold=False, space_after=0, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p


center("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ", size=12, bold=True)
center("РОССИЙСКОЙ ФЕДЕРАЦИИ", size=12, bold=True, space_after=6)
center("Федеральное государственное автономное образовательное учреждение", size=11)
center("высшего образования", size=11, space_after=6)
center("«САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ", size=12, bold=True)
center("АЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ»", size=12, bold=True, space_after=6)
center("Кафедра № 5", size=12, space_after=120)

center("ОТВЕТЫ К ЭКЗАМЕНУ", size=18, bold=True, space_after=12)
center("по дисциплине", size=14, space_after=2)
center("«Теория систем и управление технологическими изменениями»", size=14, bold=True, space_after=200)

# блок преподаватель/студент — выровнен вправо
def right_block(label, value):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_font(r, size=12)
    r2 = p.add_run(value)
    set_font(r2, size=12, bold=True)

right_block("Преподаватель: ", "Щеников Я. А.")
right_block("Выполнил студент гр. М558М: ", "Рослов Д.")
doc.add_paragraph()
doc.add_paragraph()
center("Санкт-Петербург", size=12)
center("2026", size=12)

doc.add_page_break()

# ---------- ВВОДНЫЙ АБЗАЦ ----------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Ответы на контрольные вопросы по главам учебного пособия")
set_font(r, size=14, bold=True)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(12)
r2 = p2.add_run("Щеников Я. А. Теория систем и управление технологическими изменениями: учебное пособие. — СПб., 2026")
set_font(r2, size=12, italic=True)


# ---------- ГЛАВЫ ----------
def render_md(md_text):
    lines = md_text.splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if re.fullmatch(r"[-–—*=_]{2,}", line.strip()):
            continue
        if line.startswith("# "):
            doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            add_runs(p, line[2:].strip(), size=15, base_bold=True)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            add_runs(p, line[3:].strip(), size=14, base_bold=True)
        elif re.match(r"^\s*[-–•*]\s+", line):
            txt = re.sub(r"^\s*[-–•*]\s+", "", line)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            rb = p.add_run("•  ")
            set_font(rb, size=14)
            add_runs(p, txt, size=14)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            add_runs(p, line.strip(), size=14)


for n in range(1, 8):
    md = (ANSW / f"gl{n}.md").read_text(encoding="utf-8")
    render_md(md)

doc.save(OUT)
print("Saved:", OUT)
