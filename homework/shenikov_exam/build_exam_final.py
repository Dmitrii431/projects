#!/usr/bin/env python3
"""Целевой документ: ответы строго по экзаменационным вопросам с экрана
(8 вопросов гл.5 + 2 вопроса гл.7 пособия Щеникова). Точная нумерация билета."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = Path("/home/sergei/dima-projects/homework/shenikov_exam")
ANSW = BASE / "answers"
OUT = BASE / "Экзамен_Щеников_вопросы_билета_Рослов.docx"
FONT = "Times New Roman"


def parse(md_path):
    """Возвращает {номер: текст_ответа} по ## заголовкам."""
    txt = md_path.read_text(encoding="utf-8")
    blocks = {}
    cur = None
    buf = []
    for line in txt.splitlines():
        m = re.match(r"^##\s+(\d+)\.", line)
        if m:
            if cur is not None:
                blocks[cur] = "\n".join(buf).strip()
            cur = int(m.group(1))
            buf = []
        elif line.startswith("# "):
            continue
        else:
            if cur is not None:
                buf.append(line)
    if cur is not None:
        blocks[cur] = "\n".join(buf).strip()
    return blocks


g4 = parse(ANSW / "gl4.md")
g5 = parse(ANSW / "gl5.md")
g6 = parse(ANSW / "gl6.md")
g6e = parse(ANSW / "gl6_extra.md")
g7 = parse(ANSW / "gl7.md")

# Объединённый ответ для вопроса «управление качеством И энергоэффективность»
q_qual_energy = g5[2] + "\n\n" + g5[5]

# Глава 4 — 10 вопросов с экрана (совпадают с пособием)
PART4 = [
    ("С какой целью (целями) разрабатываются цифровые двойники?", g4[1]),
    ("Какие предпосылки лежат в основе киберфизических систем?", g4[2]),
    ("В чём заключается отличие промышленного Интернета вещей от обычного Интернета вещей?", g4[3]),
    ("Объясните своими словами, как работает киберфизическая система?", g4[4]),
    ("Концептуальная модель киберфизической системы", g4[5]),
    ("Какие возможности даёт разработка и внедрение Интернета вещей?", g4[6]),
    ("Из каких составляющих состоит Интернет вещей?", g4[7]),
    ("Какие возможности даёт предприятию разработка цифровых двойников?", g4[8]),
    ("Недостатки киберфизических систем", g4[9]),
    ("Какие преимущества даёт возможность подключения всех устройств в глобальную беспроводную сеть с помощью единого стандарта?", g4[10]),
]
# Билет: глава 5 — 8 вопросов с экрана (точные формулировки)
PART5 = [
    ("Приведите информационные и термодинамические основы моделирования киберфизических систем", g5[1]),
    ("Охарактеризуйте управление качеством и энергоэффективность сложных киберфизических систем", q_qual_energy),
    ("Приведите цели и задачи управления в киберфизических системах", g5[3]),
    ("Информационные и термодинамические основы моделирования киберфизических систем", g5[1]),
    ("Охарактеризуйте управление качеством и энергоэффективность сложных киберфизических систем", q_qual_energy),
    ("Термодинамическая и информационная энтропии киберфизической системы", g5[4]),
    ("Энергоэффективность как цель киберфизической системы", g5[5]),
    ("Проблемы моделирования киберфизических систем", g5[6]),
]
# Глава 6 — 10 вопросов с экрана (1-8 из пособия, 9-10 расширены)
PART6 = [
    ("Назовите несколько отличий коллаборативного робота от классического промышленного робота", g6[1]),
    ("С помощью каких технических систем робот может распознавать действия находящегося рядом человека?", g6[2]),
    ("Что такое «обучающая выборка»? Как её сформировать и использовать для обучения искусственной нейронной сети?", g6[3]),
    ("Какие физические величины измеряют датчики, установленные на производственном оборудовании и использующиеся для предиктивного обслуживания?", g6[4]),
    ("Объясните своими словами принцип работы системы мониторинга технического состояния оборудования с использованием искусственной нейронной сети", g6[5]),
    ("Какие косвенные признаки можно наблюдать при выходе производственного оборудования из строя?", g6[6]),
    ("Приведите инструментарии проектирования и производства киберфизических систем", g6[7]),
    ("В каких случаях целесообразно использование коллаборативного робота, совместно работающего с человеком, а в каких случаях обычного промышленного робота?", g6[8]),
    ("Почему для совместной работы с человеком подходит коллаборативный робот, а не обычный?", g6e[9]),
    ("Какие технологии будут лежать в основе Индустрии 6.0?", g6e[10]),
]
# Глава 7 — 2 вопроса
PART7 = [
    ("Ключевые критерии внедрения технологических инноваций", g7[1]),
    ("Проблемы моделирования киберфизических систем", g7[2]),
]

doc = Document()
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(14)
st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
st.paragraph_format.line_spacing = 1.5
st.paragraph_format.space_after = Pt(0)
sec = doc.sections[0]
sec.left_margin, sec.right_margin = Cm(3), Cm(1.5)
sec.top_margin, sec.bottom_margin = Cm(2), Cm(2)


def setf(run, size=14, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rf)
    rf.set(qn("w:eastAsia"), FONT)


def add_runs(par, text, size=14, base_bold=False):
    for i, ch in enumerate(re.split(r"\*\*", text)):
        if ch == "":
            continue
        setf(par.add_run(ch), size=size, bold=base_bold or (i % 2 == 1))


def center(text, size=14, bold=False, sa=0, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(sa)
    setf(p.add_run(text), size=size, bold=bold, italic=italic)


# Титул
center("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ", 12, True)
center("РОССИЙСКОЙ ФЕДЕРАЦИИ", 12, True, sa=6)
center("Федеральное государственное автономное образовательное учреждение", 11)
center("высшего образования", 11, sa=6)
center("«САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ", 12, True)
center("АЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ»", 12, True, sa=6)
center("Кафедра № 5", 12, sa=110)
center("ОТВЕТЫ НА ВОПРОСЫ ЭКЗАМЕНА", 18, True, sa=12)
center("по дисциплине", 14, sa=2)
center("«Теория систем и управление технологическими изменениями»", 14, True, sa=180)


def right_block(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.space_after = Pt(2)
    setf(p.add_run(label), 12)
    setf(p.add_run(value), 12, bold=True)


right_block("Преподаватель: ", "Щеников Я. А.")
right_block("Выполнил студент гр. М558М: ", "Рослов Д.")
doc.add_paragraph()
center("Санкт-Петербург", 12)
center("2026", 12)
doc.add_page_break()


def body_par(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if re.match(r"^\s*[-–•*]\s+", text):
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        setf(p.add_run("•  "), size)
        add_runs(p, re.sub(r"^\s*[-–•*]\s+", "", text), size)
    else:
        p.paragraph_format.first_line_indent = Cm(1.25)
        add_runs(p, text.strip(), size)


def render_question(num, title, answer):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    setf(h.add_run(f"{num}. {title}"), 14, bold=True)
    for line in answer.splitlines():
        if not line.strip() or re.fullmatch(r"[-–—*=_]{2,}", line.strip()):
            continue
        body_par(line, 14)


# Часть 1 — глава 4
sec_h4 = doc.add_paragraph()
sec_h4.alignment = WD_ALIGN_PARAGRAPH.CENTER
sec_h4.paragraph_format.space_after = Pt(8)
setf(sec_h4.add_run("ВОПРОСЫ К ЭКЗАМЕНУ (Глава 4. Киберфизические системы)"), 13, bold=True)
for i, (title, ans) in enumerate(PART4, 1):
    render_question(i, title, ans)

doc.add_page_break()
sec_h = doc.add_paragraph()
sec_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
sec_h.paragraph_format.space_after = Pt(8)
setf(sec_h.add_run("ВОПРОСЫ К ЭКЗАМЕНУ (Глава 5. Цели и задачи управления в КФС)"), 13, bold=True)
for i, (title, ans) in enumerate(PART5, 1):
    render_question(i, title, ans)

doc.add_page_break()
sec_h6 = doc.add_paragraph()
sec_h6.alignment = WD_ALIGN_PARAGRAPH.CENTER
sec_h6.paragraph_format.space_after = Pt(8)
setf(sec_h6.add_run("ВОПРОСЫ К ЭКЗАМЕНУ (Глава 6. Жизненный цикл КФС)"), 13, bold=True)
for i, (title, ans) in enumerate(PART6, 1):
    render_question(i, title, ans)

doc.add_page_break()
sec_h2 = doc.add_paragraph()
sec_h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sec_h2.paragraph_format.space_after = Pt(8)
setf(sec_h2.add_run("ВОПРОСЫ К ЭКЗАМЕНУ (Глава 7. Моделирование физических процессов в КФС)"), 13, bold=True)
for i, (title, ans) in enumerate(PART7, 1):
    render_question(i, title, ans)

doc.save(OUT)
print("Saved:", OUT)
