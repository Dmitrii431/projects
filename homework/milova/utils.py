"""Утилиты для сборки отчётов по Миловой."""
import os
import shutil
import subprocess
import tempfile

from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from PIL import Image

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Mm, Inches


STUDENT_NAME_FULL = "Рослов Дмитрий Сергеевич"
STUDENT_INITIALS = "Д. С. Рослов"
GROUP = "М558М"
YEAR = "2026"
TEACHER = "В. М. Милова"
TEACHER_DEGREE = "канд. техн. наук, доцент"
DISCIPLINE = "Управление качеством организационных систем"


LO_PROFILE = "/home/sergei/.lo_profile_milova"


def _last_cell(ws):
    max_row, max_col = 0, 0
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is not None and str(cell.value).strip() != "":
                if cell.row > max_row:
                    max_row = cell.row
                if cell.column > max_col:
                    max_col = cell.column
    return max_row, max_col


def xlsx_sheet_to_png(xlsx_path: str, sheet_name: str, out_png: str, dpi: int = 150):
    """Конвертирует один лист xlsx в PNG (с обрезкой белых полей)."""
    with tempfile.TemporaryDirectory() as td:
        # data_only=True — openpyxl возвращает закэшированные значения вместо
        # формул, чтобы LibreOffice (без Java) не приходилось пересчитывать.
        wb = load_workbook(xlsx_path, data_only=True)
        for sh in list(wb.sheetnames):
            if sh != sheet_name:
                del wb[sh]
        ws = wb[sheet_name]
        last_row, last_col = _last_cell(ws)
        ws.print_area = f"A1:{get_column_letter(last_col)}{last_row}"
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 1
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.page_margins.left = 0.2
        ws.page_margins.right = 0.2
        ws.page_margins.top = 0.2
        ws.page_margins.bottom = 0.2

        trimmed_xlsx = os.path.join(td, "trimmed.xlsx")
        wb.save(trimmed_xlsx)

        subprocess.run(
            [
                "soffice",
                "--headless",
                f"-env:UserInstallation=file://{LO_PROFILE}",
                "--convert-to",
                "pdf",
                "--outdir",
                td,
                trimmed_xlsx,
            ],
            check=True,
            capture_output=True,
        )
        pdf_path = os.path.join(td, "trimmed.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError(f"PDF не создан для {sheet_name} в {xlsx_path}")

        prefix = os.path.join(td, "page")
        subprocess.run(
            ["pdftoppm", "-r", str(dpi), "-png", pdf_path, prefix],
            check=True,
            capture_output=True,
        )
        pages = sorted(p for p in os.listdir(td) if p.startswith("page-") and p.endswith(".png"))
        if not pages:
            raise RuntimeError(f"PNG не создан для {sheet_name}")

        imgs = [Image.open(os.path.join(td, p)).convert("RGB") for p in pages]

        # Обрезаем каждую страницу по содержимому.
        cropped = []
        for img in imgs:
            inv = Image.eval(img, lambda x: 255 - x)
            bbox = inv.getbbox()
            if bbox:
                pad = 20
                bbox = (
                    max(0, bbox[0] - pad),
                    max(0, bbox[1] - pad),
                    min(img.size[0], bbox[2] + pad),
                    min(img.size[1], bbox[3] + pad),
                )
                cropped.append(img.crop(bbox))
            else:
                cropped.append(img)

        if len(cropped) == 1:
            result = cropped[0]
        else:
            # Склеиваем страницы вертикально.
            max_w = max(im.width for im in cropped)
            total_h = sum(im.height for im in cropped) + 20 * (len(cropped) - 1)
            result = Image.new("RGB", (max_w, total_h), "white")
            y = 0
            for im in cropped:
                result.paste(im, ((max_w - im.width) // 2, y))
                y += im.height + 20

        os.makedirs(os.path.dirname(out_png), exist_ok=True)
        result.save(out_png, "PNG", optimize=True)
        return out_png


def set_cell_border(cell, **kwargs):
    """Устанавливает границы ячейки таблицы. kwargs: top/bottom/start/end → dict{sz,val,color}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "start", "bottom", "end"):
        if edge in kwargs:
            attrs = kwargs[edge]
            tag = OxmlElement(f"w:{edge}")
            for k, v in attrs.items():
                tag.set(qn(f"w:{k}"), str(v))
            existing = tcBorders.find(qn(f"w:{edge}"))
            if existing is not None:
                tcBorders.remove(existing)
            tcBorders.append(tag)


def add_all_borders(table):
    border = {"sz": 4, "val": "single", "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)


def set_font(run, name="Times New Roman", size=14, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def configure_document(doc: Document):
    """Настраивает страницу и стиль документа."""
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def add_paragraph(doc, text, *, bold=False, italic=False, align=None, size=14, space_after=0, first_line_indent=None, line_spacing=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_title_page(doc: Document, pr_number: int, pr_title: str):
    configure_document(doc)

    for line in (
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        "федеральное государственное автономное образовательное учреждение высшего образования",
        "«САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ АЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ»",
    ):
        add_paragraph(doc, line, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, line_spacing=1.15)
    add_paragraph(doc, "", size=12)
    add_paragraph(doc, "КАФЕДРА №5", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, line_spacing=1.15)

    for _ in range(2):
        add_paragraph(doc, "", size=12)

    # Строка «ОТЧЕТ ЗАЩИЩЕН С ОЦЕНКОЙ»
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ОТЧЕТ ")
    set_font(r, size=12, bold=True)
    r = p.add_run("\t\t\t\t")
    set_font(r, size=12)
    r = p.add_run("ЗАЩИЩЕН С ОЦЕНКОЙ")
    set_font(r, size=12, bold=True)

    add_paragraph(doc, "ПРЕПОДАВАТЕЛЬ", bold=True, size=12, line_spacing=1.15)

    # Таблица преподавателя: 2 строки × 5 колонок.
    t = doc.add_table(rows=2, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [Cm(5), Cm(2.5), Cm(3.5), Cm(0.3), Cm(5)]
    for i, w in enumerate(col_widths):
        for row in t.rows:
            row.cells[i].width = w
    top = {"sz": 4, "val": "single", "color": "000000"}
    c00, c01, c02, c03, c04 = t.rows[0].cells
    c10, c11, c12, c13, c14 = t.rows[1].cells
    # Верхний ряд: учёная степень | | | | ФИО
    for cell, text in ((c00, TEACHER_DEGREE), (c04, TEACHER)):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=12)
    # Нижний ряд — подписи под линиями.
    for cell, text in (
        (c10, "должность, уч. степень, звание"),
        (c12, "подпись, дата"),
        (c14, "инициалы, фамилия"),
    ):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=10, italic=True)
    # Линии: верхняя граница нижнего ряда (под подписью) только у c10, c12, c14.
    for cell in (c10, c12, c14):
        set_cell_border(cell, top=top)

    for _ in range(3):
        add_paragraph(doc, "", size=12)

    # Основной блок отчёта.
    for line in (
        f"ОТЧЕТ ПО ПРАКТИЧЕСКОЙ РАБОТЕ №{pr_number}",
        pr_title.upper(),
    ):
        add_paragraph(doc, line, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)
    add_paragraph(
        doc,
        f"по дисциплине: {DISCIPLINE}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
        line_spacing=1.15,
    )

    for _ in range(3):
        add_paragraph(doc, "", size=12)

    add_paragraph(doc, "РАБОТУ ВЫПОЛНИЛ", bold=True, size=12, line_spacing=1.15)

    # Таблица студента.
    t = doc.add_table(rows=2, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [Cm(3.5), Cm(1.8), Cm(1), Cm(3.5), Cm(0.3), Cm(5)]
    for i, w in enumerate(col_widths):
        for row in t.rows:
            row.cells[i].width = w
    c00, c01, c02, c03, c04, c05 = t.rows[0].cells
    c10, c11, c12, c13, c14, c15 = t.rows[1].cells
    for cell, text in (
        (c00, "СТУДЕНТ ГР. №"),
        (c01, GROUP),
        (c05, STUDENT_INITIALS),
    ):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if cell is not c00 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=12, bold=(cell is c00))
    for cell, text in ((c13, "подпись, дата"), (c15, "инициалы, фамилия")):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=10, italic=True)
    # Линия под подписью группы (c01) и двумя подписями справа.
    for cell in (c11, c13, c15):
        set_cell_border(cell, top=top)

    # Пустые строки до низа.
    for _ in range(6):
        add_paragraph(doc, "", size=12)

    add_paragraph(
        doc,
        f"Санкт-Петербург {YEAR}",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
        line_spacing=1.15,
    )

    doc.add_page_break()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    size = 16 if level == 1 else 14
    set_font(r, size=size, bold=True)
    return p


def add_body(doc, text):
    return add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        size=14,
        first_line_indent=Cm(1.25),
        line_spacing=1.5,
    )


def add_figure(doc, png_path, caption, max_width_cm=16.0):
    """Вставляет картинку по центру + подпись «Рисунок N — …»."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    r.add_picture(png_path, width=Cm(max_width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.space_before = Pt(0)
    rc = cap.add_run(caption)
    set_font(rc, size=12, italic=True)


def add_data_table(doc, headers, rows, *, first_col_bold=True, align_center=True):
    """Word-таблица: первая строка — заголовки, далее данные."""
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Заголовки
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(h))
        set_font(r, size=12, bold=True)
    # Данные
    for ri, row in enumerate(rows, start=1):
        for i, v in enumerate(row):
            cell = t.rows[ri].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run("" if v is None else str(v))
            set_font(r, size=12, bold=(first_col_bold and i == 0))
    add_all_borders(t)
    return t
