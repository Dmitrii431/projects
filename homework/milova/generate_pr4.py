"""Генератор отчёта по ПР4 — Принятие решений в условиях неопределённости и риска."""
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from PIL import Image
import subprocess

from utils import (
    LO_PROFILE,
    add_body,
    add_data_table,
    add_figure,
    add_heading,
    add_paragraph,
    add_title_page,
    xlsx_sheet_to_png,
)


XLSX = "/tmp/ПР4_Принятие_решений_Рослов-BQACAgIAAyEFAATdbEHnAAK6kWnlIM2FUB6-zvSS5g1BWrNtQh_sAALppgACIYUoSwscV5CR72ayOwQ.xlsx"
OUT_DIR = "/home/sergei/dima-projects/homework/milova"
IMG_DIR = os.path.join(OUT_DIR, "images_pr4")


def _crop_rows(xlsx_path: str, sheet: str, first_row: int, last_row: int, out_png: str, dpi=150):
    """Конвертирует диапазон строк листа в PNG."""
    with tempfile.TemporaryDirectory() as td:
        wb = load_workbook(xlsx_path, data_only=True)
        for sh in list(wb.sheetnames):
            if sh != sheet:
                del wb[sh]
        ws = wb[sheet]
        max_col = 0
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None and str(cell.value).strip():
                    if cell.column > max_col:
                        max_col = cell.column
        ws.print_area = f"A{first_row}:{get_column_letter(max_col)}{last_row}"
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 1
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.page_margins.left = 0.2
        ws.page_margins.right = 0.2
        ws.page_margins.top = 0.2
        ws.page_margins.bottom = 0.2

        trimmed = os.path.join(td, "trimmed.xlsx")
        wb.save(trimmed)
        subprocess.run(
            [
                "soffice",
                "--headless",
                f"-env:UserInstallation=file://{LO_PROFILE}",
                "--convert-to",
                "pdf",
                "--outdir",
                td,
                trimmed,
            ],
            check=True,
            capture_output=True,
        )
        prefix = os.path.join(td, "page")
        subprocess.run(["pdftoppm", "-r", str(dpi), "-png", os.path.join(td, "trimmed.pdf"), prefix], check=True, capture_output=True)
        pages = sorted(p for p in os.listdir(td) if p.startswith("page-") and p.endswith(".png"))
        imgs = [Image.open(os.path.join(td, p)).convert("RGB") for p in pages]
        cropped = []
        for img in imgs:
            inv = Image.eval(img, lambda x: 255 - x)
            bbox = inv.getbbox()
            if bbox:
                pad = 20
                bbox = (
                    max(0, bbox[0] - pad),
                    max(0, bbox[1] - pad),
                    min(img.size[0], bbox[2] + pad),
                    min(img.size[1], bbox[3] + pad),
                )
                cropped.append(img.crop(bbox))
            else:
                cropped.append(img)
        if len(cropped) == 1:
            result = cropped[0]
        else:
            max_w = max(im.width for im in cropped)
            total_h = sum(im.height for im in cropped) + 20 * (len(cropped) - 1)
            result = Image.new("RGB", (max_w, total_h), "white")
            y = 0
            for im in cropped:
                result.paste(im, ((max_w - im.width) // 2, y))
                y += im.height + 20
        os.makedirs(os.path.dirname(out_png), exist_ok=True)
        result.save(out_png, "PNG", optimize=True)
        return out_png


def main():
    os.makedirs(IMG_DIR, exist_ok=True)
    doc = Document()
    add_title_page(doc, pr_number=4, pr_title="Принятие решений в условиях неопределённости и риска")

    add_heading(doc, "Цель работы", level=2)
    add_body(
        doc,
        "Ознакомиться с классическими критериями принятия решений в условиях "
        "неопределённости (Лапласа, Вальда, Сэвиджа, Гурвица) и в условиях риска "
        "(математическое ожидание), применить их к заданной платёжной матрице и "
        "сравнить получаемые оптимальные стратегии.",
    )

    add_heading(doc, "Краткие теоретические сведения", level=2)
    add_body(
        doc,
        "Задача принятия решений в условиях неопределённости задаётся платёжной матрицей "
        "A размера m × n, где a_ij — выигрыш при выборе стратегии A_i и реализации "
        "состояния природы S_j. В зависимости от имеющейся информации о вероятностях "
        "состояний и отношения ЛПР к риску применяются различные критерии выбора "
        "оптимальной стратегии.",
    )
    add_body(
        doc,
        "Критерий Лапласа основан на принципе недостаточного основания и предполагает "
        "равновероятность исходов: L(i) = (1/n)·Σ a_ij → max. Критерий Вальда "
        "(максимин) — наиболее осторожный: W(i) = min_j a_ij → max. Критерий Сэвиджа "
        "минимизирует максимальное сожаление: сначала строится матрица сожалений "
        "r_ij = max_k a_kj − a_ij, затем выбирается стратегия с минимальным max r_ij. "
        "Критерий Гурвица — промежуточный между оптимизмом и пессимизмом: "
        "G(i) = α·min a_ij + (1−α)·max a_ij → max. При известных вероятностях y_j "
        "состояний природы применяется критерий математического ожидания "
        "M(i) = Σ a_ij·y_j → max.",
    )

    add_heading(doc, "Исходные данные", level=2)
    add_body(
        doc,
        "Задана платёжная матрица из 4 стратегий (A1…A4) и 4 исходов (S1…S4):",
    )
    add_data_table(
        doc,
        ["Стратегии \\ Исходы", "S1", "S2", "S3", "S4"],
        [
            ["A1", 10, 11, 12, 13],
            ["A2", 5, 5, 5, 25],
            ["A3", 5, 20, 20, 20],
            ["A4", 20, 5, 5, 5],
        ],
    )
    add_paragraph(doc, "", size=6)
    add_body(
        doc,
        "Для критерия Гурвица рассмотрены два значения параметра оптимизма: α = 0,2 "
        "(ближе к оптимистической оценке) и α = 0,8 (ближе к пессимистической). "
        "Для критерия математического ожидания заданы вероятности исходов "
        "y = (0,6; 0,1; 0,2; 0,1).",
    )

    add_heading(doc, "Решение", level=2)
    add_body(
        doc,
        "Расчёты по всем критериям выполнены в MS Excel на одном листе. Для каждого "
        "критерия построен отдельный блок, содержащий формулу, вспомогательные "
        "столбцы (среднее, минимум, максимум, матрица сожалений, произведения на "
        "вероятности) и выделение оптимальной стратегии. Ниже приведены рисунки с "
        "расчётом по каждому критерию.",
    )

    fig = 1
    # Блок 1 — Лаплас (строки 10–17)
    png = os.path.join(IMG_DIR, "laplace.png")
    _crop_rows(XLSX, "ПР4", 10, 17, png)
    add_figure(doc, png, f"Рисунок {fig} — Критерий Лапласа: расчёт среднего выигрыша", max_width_cm=15.5)
    fig += 1
    add_body(
        doc,
        "По критерию Лапласа L(A1) = 11,5; L(A2) = 10; L(A3) = 16,25; L(A4) = 8,75. "
        "Максимальное среднее значение даёт стратегия A3 (L* = 16,25).",
    )

    # Блок 2 — Вальд (строки 19–26)
    png = os.path.join(IMG_DIR, "wald.png")
    _crop_rows(XLSX, "ПР4", 19, 26, png)
    add_figure(doc, png, f"Рисунок {fig} — Критерий Вальда: максимин", max_width_cm=15.5)
    fig += 1
    add_body(
        doc,
        "Критерий Вальда (максимин): W(A1) = 10; W(A2) = W(A3) = W(A4) = 5. Оптимальная "
        "стратегия A1 с гарантированным выигрышем 10.",
    )

    # Блок 3 — Сэвидж (строки 28–36)
    png = os.path.join(IMG_DIR, "savage.png")
    _crop_rows(XLSX, "ПР4", 28, 36, png)
    add_figure(doc, png, f"Рисунок {fig} — Критерий Сэвиджа: матрица сожалений", max_width_cm=15.5)
    fig += 1
    add_body(
        doc,
        "Построена матрица сожалений по формуле r_ij = max_k a_kj − a_ij. Максимальные "
        "сожаления по стратегиям: A1 — 12, A2 — 15, A3 — 15, A4 — 20. Минимум достигается "
        "на A1 (C = 12).",
    )

    # Блок 4 — Гурвиц α=0.2 (строки 38–46)
    png = os.path.join(IMG_DIR, "hurwicz_02.png")
    _crop_rows(XLSX, "ПР4", 38, 46, png)
    add_figure(doc, png, f"Рисунок {fig} — Критерий Гурвица при α = 0,2", max_width_cm=15.5)
    fig += 1
    add_body(
        doc,
        "При α = 0,2 (оптимизм): G(A1) = 12,4; G(A2) = 21; G(A3) = G(A4) = 17. Оптимальна A2.",
    )

    # Блок 5 — Гурвиц α=0.8 (строки 48–55)
    png = os.path.join(IMG_DIR, "hurwicz_08.png")
    _crop_rows(XLSX, "ПР4", 48, 55, png)
    add_figure(doc, png, f"Рисунок {fig} — Критерий Гурвица при α = 0,8", max_width_cm=15.5)
    fig += 1
    add_body(
        doc,
        "При α = 0,8 (пессимизм): G(A1) = 10,6; G(A2) = 9; G(A3) = G(A4) = 8. Оптимальна A1.",
    )

    # Блок 6 — Мат. ожидание (строки 57–65)
    png = os.path.join(IMG_DIR, "expected.png")
    _crop_rows(XLSX, "ПР4", 57, 65, png)
    add_figure(doc, png, f"Рисунок {fig} — Критерий математического ожидания", max_width_cm=15.5)
    fig += 1
    add_body(
        doc,
        "По критерию математического ожидания: M(A1) = 10,8; M(A2) = 7; M(A3) = 11; "
        "M(A4) = 14. Максимум на A4.",
    )

    # Сводная таблица.
    add_heading(doc, "Сводные результаты", level=2)
    add_data_table(
        doc,
        ["Критерий", "Формула", "Оптимальная стратегия", "Значение"],
        [
            ["Лаплас", "L = (1/n)·Σ a_ij → max", "A3", "16,25"],
            ["Вальд", "W = min a_ij → max", "A1", "10"],
            ["Сэвидж", "C = min(max r_ij)", "A1", "12"],
            ["Гурвиц (α=0,2)", "G = 0,2·min + 0,8·max", "A2", "21,0"],
            ["Гурвиц (α=0,8)", "G = 0,8·min + 0,2·max", "A1", "10,6"],
            ["Мат. ожидание", "M = Σ a_ij·y_j → max", "A4", "14,0"],
        ],
    )
    add_paragraph(doc, "", size=6)

    add_heading(doc, "Общий вывод по работе", level=2)
    add_body(
        doc,
        "В ходе работы освоены шесть классических критериев принятия решений в условиях "
        "неопределённости и риска. Разные критерии дают разные рекомендации по выбору "
        "оптимальной стратегии, что объясняется различными предпосылками о поведении "
        "среды и отношении ЛПР к риску.",
    )
    add_body(
        doc,
        "Осторожный игрок, не располагающий информацией о вероятностях состояний, "
        "выберет стратегию A1 (рекомендации критериев Вальда и Сэвиджа, а также "
        "Гурвица при α = 0,8). Оптимистически настроенный ЛПР выберет A2 (Гурвиц при "
        "α = 0,2). При известном распределении вероятностей состояний природы "
        "оптимальной становится стратегия A4 (критерий математического ожидания). "
        "Стратегия A3 оптимальна при предположении о равновероятности исходов "
        "(критерий Лапласа).",
    )
    add_body(
        doc,
        "Таким образом, выбор стратегии существенно зависит от доступной информации "
        "о состоянии среды и от уровня приемлемого риска. На практике целесообразно "
        "применять несколько критериев и принимать окончательное решение на основе "
        "их совместного анализа.",
    )

    out = os.path.join(OUT_DIR, "Отчёт_ПР4_Принятие_решений_Рослов.docx")
    doc.save(out)
    print("OK:", out)


if __name__ == "__main__":
    main()
