"""Сборка отчёта ПР3 по Вершининой — нечёткий восходящий вывод."""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

from compute_pr3 import (
    X, Y, A, B, rules, build_relation, union,
    max_min_composition, centroid,
)


STUDENT_INITIALS = "Рослов Д.С."
GROUP = "М558М"
TEACHER = "Вершинина Л.П."
DISCIPLINE = "Обработка нечёткой информации\nв системах поддержки принятия решений"
YEAR = "2026"
OUT = os.path.join(os.path.dirname(__file__), "PR3_Нечёткий_восходящий_вывод_Рослов.docx")


def set_font(run, name="Times New Roman", size=14, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "start", "bottom", "end"):
        if edge in kwargs:
            attrs = kwargs[edge]
            tag = OxmlElement(f"w:{edge}")
            for k, v in attrs.items():
                tag.set(qn(f"w:{k}"), str(v))
            existing = tcBorders.find(qn(f"w:{edge}"))
            if existing is not None:
                tcBorders.remove(existing)
            tcBorders.append(tag)


def add_all_borders(table):
    border = {"sz": 4, "val": "single", "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)


def configure(doc):
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def add_par(doc, text, *, bold=False, italic=False, align=None, size=14,
            first_line_indent=None, line_spacing=1.5, space_after=0, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_body(doc, text):
    return add_par(
        doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25), line_spacing=1.5,
    )


def add_heading(doc, text, level=1):
    size = 16 if level == 1 else 14
    return add_par(
        doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
        size=size, line_spacing=1.5, space_before=6, space_after=6,
    )


def add_title_page(doc):
    configure(doc)
    add_par(doc, "САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ\nАЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ",
            bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, "Кафедра № 5\n«Инноватики и интегрированных систем качества»",
            align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)

    for _ in range(6):
        add_par(doc, "", size=12)

    add_par(doc, "ПРАКТИЧЕСКАЯ РАБОТА № 3",
            bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, "Принятие решений на основе нечёткого\nвосходящего вывода",
            bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, f"по дисциплине «{DISCIPLINE}»",
            italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)

    for _ in range(5):
        add_par(doc, "", size=12)

    add_par(doc, f"Выполнил: студент гр. {GROUP}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)
    add_par(doc, STUDENT_INITIALS, align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, "Преподаватель:", align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)
    add_par(doc, TEACHER, align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)

    for _ in range(4):
        add_par(doc, "", size=12)

    add_par(doc, f"Санкт-Петербург\n{YEAR}",
            align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)

    doc.add_page_break()


def fmt_mu(v):
    """Форматирование степени принадлежности с одной значащей цифрой после запятой."""
    if v == 0:
        return "0"
    s = f"{v:.2f}"
    # убираем лишние нули: 0.50 -> 0,5; 1.00 -> 1
    if s.endswith("0"):
        s = s[:-1]
    if s.endswith("0"):
        s = s[:-1]
    if s.endswith("."):
        s = s[:-1]
    return s.replace(".", ",")


def add_fuzzy_set_table(doc, caption, name_mapping, sets_dict, universe):
    """Таблица вида: строки — нечёткие множества, столбцы — элементы универсума."""
    add_par(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_before=6, space_after=3)
    t = doc.add_table(rows=1 + len(name_mapping), cols=1 + len(universe))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Заголовок
    hdr = t.rows[0].cells
    p = hdr[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Множество")
    set_font(r, size=12, bold=True)
    for j, x in enumerate(universe, start=1):
        p = hdr[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(x))
        set_font(r, size=12, bold=True)
    # Данные
    for ri, (label, key) in enumerate(name_mapping, start=1):
        cell = t.rows[ri].cells[0]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=12, bold=True)
        for j, x in enumerate(universe, start=1):
            cell = t.rows[ri].cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(fmt_mu(sets_dict[key][x]))
            set_font(r, size=12)
    add_all_borders(t)


def add_relation_table(doc, caption, M, row_labels, col_labels, row_header, col_header):
    add_par(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_before=6, space_after=3)
    t = doc.add_table(rows=2 + len(row_labels), cols=1 + len(col_labels))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Угловой блок — "row_header \ col_header"
    cell = t.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(row_header)
    set_font(r, size=12, bold=True)
    # Верхняя шапка — col_header (объединён по col_labels)
    cell_top = t.rows[0].cells[1]
    for j in range(2, 1 + len(col_labels)):
        cell_top.merge(t.rows[0].cells[j])
    p = cell_top.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(col_header)
    set_font(r, size=12, bold=True)
    # Вторая строка — значения col_labels
    cell = t.rows[1].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("")
    set_font(r, size=12)
    for j, c in enumerate(col_labels, start=1):
        cell = t.rows[1].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(c))
        set_font(r, size=12, bold=True)
    # Данные
    for ri, rlabel in enumerate(row_labels, start=2):
        cell = t.rows[ri].cells[0]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(rlabel))
        set_font(r, size=12, bold=True)
        for j, _ in enumerate(col_labels, start=1):
            cell = t.rows[ri].cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(fmt_mu(M[ri - 2][j - 1]))
            set_font(r, size=12)
    add_all_borders(t)


def main():
    doc = Document()
    add_title_page(doc)

    # Введение
    add_heading(doc, "ВВЕДЕНИЕ")
    add_body(doc,
        "Нечёткий восходящий (прямой) вывод представляет собой модель "
        "рассуждений, при которой из входного наблюдения по известной базе "
        "продукционных правил выводится заключение, выраженное нечётким "
        "множеством. Основу модели составляет аппарат нечётких отношений: "
        "каждое правило «ЕСЛИ A то B» интерпретируется как нечёткое "
        "отношение импликации на декартовом произведении универсумов, а "
        "итоговый вывод формируется максиминной композицией входного "
        "множества и объединения всех отношений.")
    add_body(doc,
        "В работе разработана модель принятия решений на основе нечёткого "
        "восходящего вывода для задачи выбора мощности искусственной "
        "подсветки рабочего места в офисе в зависимости от уровня "
        "естественной освещённости. Модель содержит пять продукционных "
        "правил и демонстрирует работу механизма max-min композиции на "
        "конкретном числовом примере.")

    # 1. Постановка
    add_heading(doc, "1. ПОСТАНОВКА ЗАДАЧИ")
    add_body(doc,
        "Требуется построить модель принятия решений на основе нечёткого "
        "восходящего вывода с использованием нечётких отношений. Для "
        "этого необходимо:")
    for line in (
        "1) определить входной и выходной универсумы X и Y;",
        "2) задать лингвистические переменные на X и Y, описав их "
        "нечёткими множествами;",
        "3) сформулировать не менее пяти продукционных правил вида "
        "«ЕСЛИ X есть A THEN Y есть B»;",
        "4) построить нечёткие отношения импликации R_k для каждого "
        "правила и общее отношение R как их объединение;",
        "5) продемонстрировать работу модели на числовом примере: "
        "при заданном наблюдении A′ выполнить max-min композицию "
        "B′ = A′ ∘ R и выполнить дефаззификацию методом центра тяжести.",
    ):
        add_body(doc, line)

    add_body(doc,
        "Нечёткое отношение импликации для правила «ЕСЛИ X есть A THEN Y "
        "есть B» строится по правилу Мамдани: μ_R(x, y) = min(μ_A(x), "
        "μ_B(y)). Объединение отношений в единую базу правил выполняется "
        "поэлементным максимумом: μ_R(x, y) = max_k μ_{R_k}(x, y). "
        "Восходящий вывод для входного нечёткого множества A′ вычисляется "
        "по правилу max-min композиции: μ_{B′}(y) = max_x min(μ_{A′}(x), "
        "μ_R(x, y)). Результирующее нечёткое множество B′ при "
        "необходимости дефаззифицируется в чёткое число методом центра "
        "тяжести.")

    # 2. Предметная область
    add_heading(doc, "2. ПРЕДМЕТНАЯ ОБЛАСТЬ И УНИВЕРСУМЫ")
    add_body(doc,
        "В качестве предметной области рассматривается задача "
        "автоматического управления искусственной подсветкой рабочего "
        "места оператора в офисном помещении. Датчик измеряет уровень "
        "естественной освещённости на поверхности стола, а система "
        "определяет рекомендуемую мощность включаемой настольной лампы. "
        "Особенность задачи — плавный характер зависимости: одним и тем "
        "же «словесным» уровням освещённости соответствуют диапазоны "
        "значений, что делает аппарат нечёткого вывода естественным "
        "инструментом моделирования.")

    add_body(doc,
        "Входной универсум — уровень естественной освещённости на "
        "рабочем месте, измеряемый в люксах: X = {100, 300, 500, 700, "
        "900} лк.")
    add_body(doc,
        "Выходной универсум — мощность искусственной подсветки, "
        "измеряемая в ваттах: Y = {10, 20, 30, 40, 50} Вт.")

    add_body(doc,
        "На входном универсуме задана лингвистическая переменная "
        "«естественная освещённость» с пятью терм-значениями: A₁ — очень "
        "низкая, A₂ — низкая, A₃ — средняя, A₄ — высокая, A₅ — очень "
        "высокая. На выходном универсуме задана лингвистическая "
        "переменная «мощность подсветки» с пятью терм-значениями: B₁ — "
        "максимальная, B₂ — высокая, B₃ — средняя, B₄ — низкая, B₅ — "
        "минимальная.")

    add_fuzzy_set_table(
        doc,
        "Таблица 1 — Функции принадлежности нечётких множеств A₁–A₅ на универсуме X (лк)",
        [
            ("A₁ очень низкая",  "A1_очень_низкая"),
            ("A₂ низкая",        "A2_низкая"),
            ("A₃ средняя",       "A3_средняя"),
            ("A₄ высокая",       "A4_высокая"),
            ("A₅ очень высокая", "A5_очень_высокая"),
        ],
        A, X,
    )

    add_fuzzy_set_table(
        doc,
        "Таблица 2 — Функции принадлежности нечётких множеств B₁–B₅ на универсуме Y (Вт)",
        [
            ("B₁ максимальная", "B1_максимальная"),
            ("B₂ высокая",      "B2_высокая"),
            ("B₃ средняя",      "B3_средняя"),
            ("B₄ низкая",       "B4_низкая"),
            ("B₅ минимальная",  "B5_минимальная"),
        ],
        B, Y,
    )

    # 3. База правил
    add_heading(doc, "3. БАЗА ПРОДУКЦИОННЫХ ПРАВИЛ")
    add_body(doc,
        "База знаний содержит пять продукционных правил, описывающих "
        "обратную зависимость «чем ниже естественное освещение, тем выше "
        "мощность искусственной подсветки»:")
    for line, (a_k, b_k) in zip(
        ("Правило 1:", "Правило 2:", "Правило 3:", "Правило 4:", "Правило 5:"),
        rules,
    ):
        a_label = {
            "A1_очень_низкая": "A₁ (очень низкая)",
            "A2_низкая": "A₂ (низкая)",
            "A3_средняя": "A₃ (средняя)",
            "A4_высокая": "A₄ (высокая)",
            "A5_очень_высокая": "A₅ (очень высокая)",
        }[a_k]
        b_label = {
            "B1_максимальная": "B₁ (максимальная)",
            "B2_высокая": "B₂ (высокая)",
            "B3_средняя": "B₃ (средняя)",
            "B4_низкая": "B₄ (низкая)",
            "B5_минимальная": "B₅ (минимальная)",
        }[b_k]
        add_par(
            doc, f"{line} ЕСЛИ X есть {a_label}, ТО Y есть {b_label}.",
            align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(1.25),
            line_spacing=1.5,
        )
    add_body(doc,
        "Правила покрывают весь диапазон входных значений и обеспечивают "
        "плавный переход управляющего воздействия, поскольку смежные "
        "термы частично пересекаются.")

    # 4. Нечёткие отношения
    add_heading(doc, "4. ПОСТРОЕНИЕ НЕЧЁТКИХ ОТНОШЕНИЙ")
    add_body(doc,
        "Для каждого правила R_k построим матрицу нечёткого отношения "
        "импликации на декартовом произведении X × Y по формуле "
        "μ_{R_k}(x, y) = min(μ_{A_k}(x), μ_{B_k}(y)).")

    Rs = []
    caption_map = {
        1: "Таблица 3 — Матрица отношения R₁ (правило 1: A₁ → B₁)",
        2: "Таблица 4 — Матрица отношения R₂ (правило 2: A₂ → B₂)",
        3: "Таблица 5 — Матрица отношения R₃ (правило 3: A₃ → B₃)",
        4: "Таблица 6 — Матрица отношения R₄ (правило 4: A₄ → B₄)",
        5: "Таблица 7 — Матрица отношения R₅ (правило 5: A₅ → B₅)",
    }
    for k, (a_k, b_k) in enumerate(rules, start=1):
        Rk = build_relation(a_k, b_k)
        Rs.append(Rk)
        add_relation_table(
            doc, caption_map[k], Rk, X, Y,
            row_header="X, лк", col_header="Y, Вт",
        )

    add_body(doc,
        "Общая база правил описывается одним нечётким отношением R, "
        "которое получается объединением отношений R₁, R₂, …, R₅: "
        "μ_R(x, y) = max_k μ_{R_k}(x, y).")

    R = union(*Rs)
    add_relation_table(
        doc,
        "Таблица 8 — Общее нечёткое отношение R = R₁ ∪ R₂ ∪ R₃ ∪ R₄ ∪ R₅",
        R, X, Y, row_header="X, лк", col_header="Y, Вт",
    )

    # 5. Пример
    add_heading(doc, "5. ПРИМЕР РАСЧЁТА")
    add_body(doc,
        "Пусть датчик зафиксировал уровень естественной освещённости "
        "около 420 лк. Это наблюдение формализуется нечётким множеством "
        "A′ на универсуме X, имеющим максимум между термами «низкая» (300 "
        "лк) и «средняя» (500 лк) и слегка смещённым в сторону меньших "
        "значений:")

    A_prime = {100: 0.1, 300: 0.8, 500: 0.7, 700: 0.2, 900: 0.0}
    add_par(doc,
        "A′ = {100/0,1; 300/0,8; 500/0,7; 700/0,2; 900/0}.",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.5,
        space_before=3, space_after=3,
    )

    add_body(doc,
        "Восходящий вывод выполняется по правилу max-min композиции "
        "B′ = A′ ∘ R. Для каждого значения y ∈ Y вычисляется "
        "μ_{B′}(y) = max_x min(μ_{A′}(x), μ_R(x, y)).")

    add_body(doc, "Подробный расчёт значений:")
    for j, y in enumerate(Y):
        parts = []
        for i, x in enumerate(X):
            parts.append(f"min({fmt_mu(A_prime[x])}; {fmt_mu(R[i][j])})")
        vals = [min(A_prime[X[i]], R[i][j]) for i in range(len(X))]
        add_par(
            doc,
            f"μ_B′({y}) = max[{'; '.join(parts)}] = {fmt_mu(max(vals))};",
            align=WD_ALIGN_PARAGRAPH.LEFT, size=12, line_spacing=1.3,
        )

    B_prime = max_min_composition(A_prime, R)
    add_body(doc, "Итоговое выходное нечёткое множество:")
    add_par(doc,
        "B′ = {10/" + fmt_mu(B_prime[10]) + "; 20/" + fmt_mu(B_prime[20]) +
        "; 30/" + fmt_mu(B_prime[30]) + "; 40/" + fmt_mu(B_prime[40]) +
        "; 50/" + fmt_mu(B_prime[50]) + "}.",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.5,
        space_before=3, space_after=3,
    )

    # Таблица B'
    add_par(doc, "Таблица 9 — Нечёткое множество B′ = A′ ∘ R на универсуме Y",
            italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12,
            space_before=6, space_after=3)
    t = doc.add_table(rows=2, cols=1 + len(Y))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    p = hdr[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Y, Вт")
    set_font(r, size=12, bold=True)
    for j, y in enumerate(Y, start=1):
        p = hdr[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(y))
        set_font(r, size=12, bold=True)
    cell = t.rows[1].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("μ_B′(y)")
    set_font(r, size=12, bold=True)
    for j, y in enumerate(Y, start=1):
        cell = t.rows[1].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(fmt_mu(B_prime[y]))
        set_font(r, size=12)
    add_all_borders(t)

    y_star = centroid(B_prime)
    num_str = " + ".join(f"{y}·{fmt_mu(B_prime[y])}" for y in Y if B_prime[y] > 0)
    den_str = " + ".join(fmt_mu(B_prime[y]) for y in Y if B_prime[y] > 0)
    num_val = sum(y * B_prime[y] for y in Y)
    den_val = sum(B_prime[y] for y in Y)

    add_body(doc,
        "Дефаззификация результата выполняется методом центра тяжести. "
        "Чёткое значение выходной переменной y* вычисляется по формуле "
        "y* = Σ y·μ_B′(y) / Σ μ_B′(y):")
    num_ru = f"{num_val:.1f}".replace(".", ",")
    den_ru = f"{den_val:.1f}".replace(".", ",")
    y_star_ru = f"{y_star:.2f}".replace(".", ",")
    formula_line = (
        f"y* = ({num_str}) / ({den_str}) = "
        f"{num_ru} / {den_ru} ≈ {y_star_ru} Вт."
    )
    add_par(doc,
        formula_line,
        align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.5,
        space_before=3, space_after=3,
    )

    add_body(doc,
        f"Таким образом, при естественной освещённости около 420 лк "
        f"модель рекомендует установить мощность искусственной подсветки "
        f"приблизительно {int(round(y_star))} Вт. Значение лежит между "
        f"термами «средняя» (30 Вт) и «высокая» (40 Вт) мощность, что "
        f"соответствует интуитивному решению: наблюдение ближе к терму "
        f"«низкая» (300 лк, μ = 0,8), поэтому система увеличивает "
        f"мощность по сравнению со средним значением.")

    # Заключение
    add_heading(doc, "ЗАКЛЮЧЕНИЕ")
    add_body(doc,
        "В работе разработана модель принятия решений на основе нечёткого "
        "восходящего вывода. Определены входной и выходной универсумы, "
        "построены лингвистические переменные с пятью терм-значениями на "
        "каждом универсуме, сформулирована база из пяти продукционных "
        "правил, каждое из которых представлено нечётким отношением "
        "импликации Мамдани. Объединением частных отношений получена "
        "обобщённая матрица R.")
    add_body(doc,
        "Работа модели проиллюстрирована на примере обработки "
        "наблюдения «естественная освещённость около 420 лк». Max-min "
        "композиция A′ ∘ R позволила получить выходное нечёткое "
        "множество B′, а метод центра тяжести дал чёткую рекомендацию "
        f"y* ≈ {int(round(y_star))} Вт.")
    add_body(doc,
        "Разработанная модель обладает прозрачной интерпретацией на "
        "уровне продукционных правил, обеспечивает плавный переход "
        "между режимами работы и может быть расширена за счёт "
        "добавления новых правил, термов или входных переменных без "
        "изменения математического аппарата.")

    # Источники
    add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    sources = [
        "1. Заде, Л. Понятие лингвистической переменной и его применение "
        "к принятию приближённых решений / Л. Заде. — М. : Мир, 1976. — 165 с.",
        "2. Борисов, А. Н. Принятие решений на основе нечётких моделей / "
        "А. Н. Борисов, О. А. Крумберг, И. П. Фёдоров. — Рига : Зинатне, "
        "1990. — 184 с.",
        "3. Кофман, А. Введение в теорию нечётких множеств / А. Кофман. — "
        "М. : Радио и связь, 1982. — 432 с.",
        "4. Леоненков, А. В. Нечёткое моделирование в среде MATLAB и "
        "fuzzyTECH / А. В. Леоненков. — СПб. : БХВ-Петербург, 2005. — 736 с.",
        "5. Круглов, В. В. Нечёткая логика и искусственные нейронные "
        "сети / В. В. Круглов, М. И. Дли, Р. Ю. Голунов. — М. : Физматлит, "
        "2001. — 224 с.",
        "6. Орлов, А. И. Теория принятия решений / А. И. Орлов. — М. : "
        "Экзамен, 2006. — 573 с.",
    ]
    for s in sources:
        add_par(doc, s, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=Cm(1.25), line_spacing=1.5)

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
