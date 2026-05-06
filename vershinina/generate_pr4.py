"""Сборка отчёта ПР4 по Вершининой — нечёткая классификационная модель
принятия решений (вариант 11 — расчёт потребления бензина)."""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

from compute_pr4 import (
    UNIVERSES, TERMS, CLASSES,
    membership, classify, classify_detailed,
    X1, X2, X3, X4, T1, T2, T3, T4,
)


STUDENT_INITIALS = "Рослов Д.С."
GROUP = "М558М"
TEACHER = "Вершинина Л.П."
DISCIPLINE = "Обработка нечёткой информации\nв системах поддержки принятия решений"
YEAR = "2026"
OUT = os.path.join(os.path.dirname(__file__), "PR4_Нечёткая_классификационная_модель_Рослов.docx")


def set_font(run, name="Times New Roman", size=14, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "start", "bottom", "end"):
        if edge in kwargs:
            attrs = kwargs[edge]
            tag = OxmlElement(f"w:{edge}")
            for k, v in attrs.items():
                tag.set(qn(f"w:{k}"), str(v))
            existing = tcBorders.find(qn(f"w:{edge}"))
            if existing is not None:
                tcBorders.remove(existing)
            tcBorders.append(tag)


def add_all_borders(table):
    border = {"sz": 4, "val": "single", "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)


def configure(doc):
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def add_par(doc, text, *, bold=False, italic=False, align=None, size=14,
            first_line_indent=None, line_spacing=1.5, space_after=0, space_before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_body(doc, text):
    return add_par(
        doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25), line_spacing=1.5,
    )


def add_heading(doc, text, level=1):
    size = 16 if level == 1 else 14
    return add_par(
        doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
        size=size, line_spacing=1.5, space_before=6, space_after=6,
    )


def add_title_page(doc):
    configure(doc)
    add_par(doc, "САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ\nАЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ",
            bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, "Кафедра № 5\n«Инноватики и интегрированных систем качества»",
            align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)

    for _ in range(6):
        add_par(doc, "", size=12)

    add_par(doc, "ПРАКТИЧЕСКАЯ РАБОТА № 4",
            bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, "Нечёткая классификационная модель\nпринятия решений",
            bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, f"по дисциплине «{DISCIPLINE}»",
            italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, "Вариант 11. Расчёт потребления бензина",
            italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)

    for _ in range(4):
        add_par(doc, "", size=12)

    add_par(doc, f"Выполнил: студент гр. {GROUP}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)
    add_par(doc, STUDENT_INITIALS, align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)
    add_par(doc, "", size=12)
    add_par(doc, "Преподаватель:", align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)
    add_par(doc, TEACHER, align=WD_ALIGN_PARAGRAPH.RIGHT, size=14, line_spacing=1.15)

    for _ in range(4):
        add_par(doc, "", size=12)

    add_par(doc, f"Санкт-Петербург\n{YEAR}",
            align=WD_ALIGN_PARAGRAPH.CENTER, size=14, line_spacing=1.15)

    doc.add_page_break()


def fmt_mu(v):
    """0.50 -> 0,5; 1.00 -> 1; 0 -> 0."""
    if v == 0:
        return "0"
    s = f"{v:.2f}"
    if s.endswith("0"):
        s = s[:-1]
    if s.endswith("0"):
        s = s[:-1]
    if s.endswith("."):
        s = s[:-1]
    return s.replace(".", ",")


def fmt_mu3(v):
    """Три знака после запятой — для итогового результата."""
    if v == 0:
        return "0"
    s = f"{v:.3f}".replace(".", ",")
    return s


def add_terms_table(doc, caption, var_name, terms_dict, universe, units):
    """Таблица функций принадлежности термов одной лингв. переменной."""
    add_par(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12,
            space_before=6, space_after=3)
    t = doc.add_table(rows=1 + len(terms_dict), cols=1 + len(universe))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    p = hdr[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"Терм \\ {var_name}, {units}")
    set_font(r, size=11, bold=True)
    for j, x in enumerate(universe, start=1):
        p = hdr[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(x))
        set_font(r, size=11, bold=True)
    for ri, (term_name, mus) in enumerate(terms_dict.items(), start=1):
        cell = t.rows[ri].cells[0]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(term_name)
        set_font(r, size=11, bold=True)
        for j, x in enumerate(universe, start=1):
            cell = t.rows[ri].cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(fmt_mu(mus[x]))
            set_font(r, size=11)
    add_all_borders(t)


def add_classes_table(doc, caption):
    """Таблица эталонных классов: строки — классы, столбцы — переменные."""
    add_par(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12,
            space_before=6, space_after=3)
    headers = ["Класс",
               "X₁ маневры",
               "X₂ стаж водителя",
               "X₃ объём двигателя",
               "X₄ состояние авто"]
    t = doc.add_table(rows=1 + len(CLASSES), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, size=11, bold=True)
    for ri, (c_key, c_def) in enumerate(CLASSES.items(), start=1):
        cells = t.rows[ri].cells
        cell_data = [
            c_def["label"],
            c_def["X1"],
            c_def["X2"],
            c_def["X3"],
            c_def["X4"],
        ]
        for j, val in enumerate(cell_data):
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            set_font(r, size=11, bold=(j == 0))
    add_all_borders(t)


def add_observation_table(doc, caption, observation):
    add_par(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12,
            space_before=6, space_after=3)
    rows_data = [
        ("X₁", "доля резких маневров", "%", observation["X1"]),
        ("X₂", "стаж водителя", "лет", observation["X2"]),
        ("X₃", "рабочий объём двигателя", "л", observation["X3"]),
        ("X₄", "износ автомобиля", "%", observation["X4"]),
    ]
    t = doc.add_table(rows=1 + len(rows_data), cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Переменная", "Описание", "Ед. изм.", "Значение"]
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, size=12, bold=True)
    for ri, (var, desc, unit, val) in enumerate(rows_data, start=1):
        for j, v in enumerate((var, desc, unit, str(val))):
            cell = t.rows[ri].cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(v)
            set_font(r, size=12, bold=(j == 0))
    add_all_borders(t)


def add_classification_table(doc, caption, observation):
    add_par(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12,
            space_before=6, space_after=3)
    headers = ["Класс", "μ(X₁)", "μ(X₂)", "μ(X₃)", "μ(X₄)", "μ_C = min"]
    rows = classify_detailed(observation)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, size=12, bold=True)
    for ri, (c_key, label, per_var) in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=11, bold=True)
        for j, var in enumerate(("X1", "X2", "X3", "X4"), start=1):
            _, mu = per_var[var]
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(fmt_mu3(mu))
            set_font(r, size=12)
        p = cells[5].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(fmt_mu3(per_var["min"]))
        set_font(r, size=12, bold=True)
    add_all_borders(t)


def main():
    doc = Document()
    add_title_page(doc)

    # Введение
    add_heading(doc, "ВВЕДЕНИЕ")
    add_body(doc,
        "Нечёткая классификационная модель принятия решений предназначена "
        "для отнесения наблюдаемого объекта к одному из заранее определённых "
        "эталонных классов в условиях, когда границы между классами размыты "
        "и значения признаков объекта известны нечётко. В отличие от "
        "решающих правил классической логики, в нечёткой модели каждое "
        "значение признака описывается лингвистической переменной с "
        "несколькими терм-значениями, а эталонный класс задаётся как "
        "комбинация термов. Принадлежность наблюдения к классу "
        "вычисляется через нечёткую конъюнкцию частных принадлежностей.")
    add_body(doc,
        "В работе разработана классификационная модель для задачи оценки "
        "уровня потребления бензина легковым автомобилем. Модель содержит "
        "четыре входные лингвистические переменные и три эталонных класса, "
        "соответствующих низкому, среднему и высокому расходу. Работа "
        "модели проиллюстрирована на конкретном числовом примере.")

    # 1. Постановка
    add_heading(doc, "1. ПОСТАНОВКА ЗАДАЧИ")
    add_body(doc,
        "В соответствии с вариантом 11 областью принятия решений является "
        "расчёт потребления бензина легковым автомобилем. На уровень "
        "расхода топлива оказывают влияние тип совершаемых маневров, "
        "уровень подготовки водителя, тип автомобиля и его техническое "
        "состояние. Все перечисленные факторы носят качественный характер "
        "и описываются естественным языком, что делает применение "
        "классических количественных моделей затруднительным.")
    add_body(doc,
        "Требуется построить нечёткую классификационную модель, "
        "позволяющую по значениям перечисленных факторов отнести "
        "автомобиль к одному из эталонных классов потребления бензина. "
        "Для построения модели необходимо:")
    for line in (
        "1) выбрать не менее четырёх лингвистических переменных, "
        "характеризующих объект классификации, и определить их "
        "универсумы;",
        "2) задать на каждом универсуме систему термов в виде нечётких "
        "множеств;",
        "3) определить не менее трёх эталонных классов, каждый из которых "
        "представляет собой комбинацию термов входных переменных;",
        "4) сформулировать алгоритм классификации, основанный на операции "
        "нечёткой конъюнкции;",
        "5) проверить работоспособность модели на конкретном числовом "
        "примере.",
    ):
        add_body(doc, line)

    add_body(doc,
        "Принадлежность наблюдаемого объекта o = (o₁, o₂, o₃, o₄) к "
        "эталонному классу C, заданному набором термов T_C¹, T_C², T_C³, "
        "T_C⁴ соответствующих лингвистических переменных, вычисляется по "
        "правилу нечёткой конъюнкции: μ_C(o) = min_i μ_{T_C^i}(o_i). "
        "Объект относится к классу C* с максимальным значением "
        "принадлежности: C* = argmax_C μ_C(o).")

    # 2. Лингвистические переменные
    add_heading(doc, "2. ЛИНГВИСТИЧЕСКИЕ ПЕРЕМЕННЫЕ")
    add_body(doc,
        "Для описания объекта классификации введены четыре "
        "лингвистические переменные.")

    add_body(doc,
        "Переменная X₁ «тип совершаемых маневров» характеризует "
        "агрессивность стиля вождения и измеряется как доля резких "
        "ускорений и торможений в общем времени движения, %. Универсум: "
        "X₁ = {0, 25, 50, 75, 100}. Терм-множество: T₁ = {спокойный, "
        "умеренный, агрессивный}.")

    add_terms_table(
        doc,
        "Таблица 1 — Функции принадлежности термов переменной X₁",
        "X₁", T1, X1, "%",
    )

    add_body(doc,
        "Переменная X₂ «уровень подготовки водителя» характеризует опыт "
        "управления автомобилем и измеряется в годах водительского "
        "стажа. Универсум: X₂ = {0, 5, 10, 20, 30}. Терм-множество: T₂ = "
        "{начинающий, средний, опытный}.")

    add_terms_table(
        doc,
        "Таблица 2 — Функции принадлежности термов переменной X₂",
        "X₂", T2, X2, "лет",
    )

    add_body(doc,
        "Переменная X₃ «тип автомобиля» характеризует размер автомобиля "
        "через рабочий объём двигателя, л. Универсум: X₃ = {1,0; 1,5; "
        "2,0; 3,0; 4,0}. Терм-множество: T₃ = {малолитражка, средний, "
        "крупный}.")

    add_terms_table(
        doc,
        "Таблица 3 — Функции принадлежности термов переменной X₃",
        "X₃", T3, X3, "л",
    )

    add_body(doc,
        "Переменная X₄ «техническое состояние автомобиля» характеризует "
        "степень износа узлов и агрегатов и измеряется в процентах "
        "выработки ресурса. Универсум: X₄ = {0, 25, 50, 75, 100}. "
        "Терм-множество: T₄ = {отличное, удовлетворительное, плохое}.")

    add_terms_table(
        doc,
        "Таблица 4 — Функции принадлежности термов переменной X₄",
        "X₄", T4, X4, "%",
    )

    # 3. Эталонные классы
    add_heading(doc, "3. ЭТАЛОННЫЕ КЛАССЫ")
    add_body(doc,
        "Объектом классификации является автомобиль вместе с его "
        "водителем, рассматриваемый как совокупность четырёх признаков. "
        "В качестве эталонных классов приняты три категории расхода "
        "бензина — низкий, средний и высокий. Каждый класс задаётся "
        "комбинацией термов входных лингвистических переменных, "
        "соответствующих типичному представителю данной категории.")

    add_classes_table(doc, "Таблица 5 — Эталонные классы")

    add_body(doc,
        "Класс C₁ «низкий расход» (5–7 л/100 км) описывает экономичную "
        "поездку: спокойные маневры, опытный водитель, малолитражный "
        "автомобиль в отличном состоянии. Класс C₂ «средний расход» (7–10 "
        "л/100 км) соответствует типичной городской эксплуатации "
        "автомобиля среднего класса водителем со средним уровнем "
        "подготовки. Класс C₃ «высокий расход» (10–14 л/100 км) "
        "соответствует наиболее затратному режиму: агрессивная езда "
        "начинающего водителя на крупном автомобиле в плохом техническом "
        "состоянии.")

    # 4. Алгоритм классификации
    add_heading(doc, "4. АЛГОРИТМ КЛАССИФИКАЦИИ")
    add_body(doc,
        "Алгоритм классификации наблюдаемого объекта o = (o₁, o₂, o₃, "
        "o₄) включает следующие шаги.")
    for line in (
        "Шаг 1. Для каждого эталонного класса C ∈ {C₁, C₂, C₃} и каждой "
        "переменной X_i (i = 1…4) определить терм T_C^i, входящий в "
        "описание класса C.",
        "Шаг 2. Вычислить степень принадлежности μ_{T_C^i}(o_i) "
        "наблюдаемого значения o_i соответствующему терму. Если значение "
        "o_i не совпадает с узлом универсума, применяется кусочно-линейная "
        "интерполяция между двумя ближайшими узлами.",
        "Шаг 3. Вычислить общую степень принадлежности объекта классу C "
        "как минимум частных принадлежностей: μ_C(o) = min_i "
        "μ_{T_C^i}(o_i).",
        "Шаг 4. Определить класс-победитель C* = argmax_C μ_C(o); "
        "значение μ_{C*}(o) рассматривается как уверенность модели в "
        "принятом решении.",
    ):
        add_body(doc, line)

    add_body(doc,
        "Использование операции min для агрегации частных "
        "принадлежностей соответствует логике конъюнкции: объект "
        "относится к классу настолько, насколько он удовлетворяет всем "
        "признакам класса одновременно. Низкое значение хотя бы по "
        "одному признаку резко снижает итоговую принадлежность, что "
        "соответствует интуитивному пониманию строгой совместимости "
        "признаков объекта с эталоном.")

    # 5. Расчёт
    add_heading(doc, "5. ПРИМЕР РАСЧЁТА")
    observation = {"X1": 35, "X2": 8, "X3": 1.6, "X4": 30}

    add_body(doc,
        "Для проверки работоспособности модели рассмотрим автомобиль "
        "молодого водителя с восьмилетним стажем. Доля резких "
        "ускорений и торможений в его поездках составляет около 35 %, "
        "автомобиль — компактный седан с двигателем 1,6 л в относительно "
        "хорошем техническом состоянии (около 30 % износа).")

    add_observation_table(
        doc, "Таблица 6 — Значения признаков наблюдаемого объекта", observation,
    )

    add_body(doc,
        "Поскольку значения o₁ = 35 %, o₃ = 1,6 л и o₄ = 30 % не "
        "совпадают с узлами соответствующих универсумов, степени "
        "принадлежности термам определяются кусочно-линейной "
        "интерполяцией. Например, для μ_спокойный(35) узлы 25 и 50 дают "
        "значения 0,6 и 0,1 соответственно, откуда μ_спокойный(35) = 0,6 "
        "+ (0,1 − 0,6) · (35 − 25) / (50 − 25) = 0,4. Аналогично "
        "вычисляются остальные степени принадлежности.")

    rows = classify_detailed(observation)
    add_body(doc,
        "Результаты расчёта степеней принадлежности наблюдаемого "
        "объекта термам для каждого эталонного класса, а также итоговые "
        "значения μ_C(o), приведены в таблице.")

    add_classification_table(
        doc, "Таблица 7 — Расчёт принадлежности объекта эталонным классам", observation,
    )

    add_body(doc, "Подробная запись расчёта по каждому классу:")
    for c_key, label, per_var in rows:
        parts = []
        for var in ("X1", "X2", "X3", "X4"):
            term, mu = per_var[var]
            parts.append(f"μ_{{{var}={term}}}({observation[var]}) = {fmt_mu3(mu)}")
        line = f"{label}: " + "; ".join(parts) + f"; μ_C = min(...) = {fmt_mu3(per_var['min'])}."
        add_par(doc, line, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=Cm(1.25), line_spacing=1.5, size=12)

    result = classify(observation)
    winner = max(result, key=result.get)
    winner_label = CLASSES[winner]["label"]

    add_body(doc,
        f"Сравнение итоговых значений показывает, что наибольшая степень "
        f"принадлежности достигается для класса {winner_label}: μ = "
        f"{fmt_mu3(result[winner])}. Следовательно, наблюдаемый объект "
        f"относится моделью к классу {winner_label}, что соответствует "
        f"средней категории потребления бензина — 7–10 л/100 км. "
        f"Полученный результат согласуется с интуитивной оценкой: "
        f"автомобиль среднего размера в умеренно нагруженном режиме "
        f"эксплуатации водителем со средним стажем должен показывать "
        f"типичный для городских условий расход топлива.")

    add_body(doc,
        "Для классов C₁ и C₃ итоговая принадлежность оказалась "
        "существенно ниже из-за низких значений отдельных частных "
        "принадлежностей. Так, для класса «низкий расход» определяющим "
        "оказался стаж водителя — 8 лет соответствуют терму «опытный» "
        "лишь со степенью 0,18; для класса «высокий расход» решающую "
        "роль сыграло несоответствие компактного автомобиля 1,6 л "
        "терму «крупный».")

    # Заключение
    add_heading(doc, "ЗАКЛЮЧЕНИЕ")
    add_body(doc,
        "В работе разработана нечёткая классификационная модель "
        "принятия решений для задачи оценки уровня потребления бензина "
        "легковым автомобилем. Введены четыре лингвистические "
        "переменные — тип совершаемых маневров, уровень подготовки "
        "водителя, рабочий объём двигателя и техническое состояние "
        "автомобиля; на каждом универсуме построена система из трёх "
        "термов с заданными функциями принадлежности.")
    add_body(doc,
        "В качестве эталонных классов выбраны три категории расхода "
        "топлива — низкий, средний и высокий, представленные "
        "комбинациями термов входных переменных. Сформулирован алгоритм "
        "классификации, основанный на операции нечёткой конъюнкции: "
        "μ_C(o) = min_i μ_{T_C^i}(o_i).")
    add_body(doc,
        f"На числовом примере (автомобиль 1,6 л, водитель со стажем "
        f"8 лет, 35 % резких маневров, 30 % износа) показано, что "
        f"наблюдаемый объект относится моделью к классу «{winner_label}» "
        f"со степенью принадлежности {fmt_mu3(result[winner])}. "
        f"Полученный результат согласуется с интуитивной оценкой "
        f"экспертов и подтверждает работоспособность построенной модели.")
    add_body(doc,
        "Разработанная классификационная модель обладает прозрачной "
        "интерпретацией на уровне эталонных классов, легко расширяется "
        "за счёт добавления новых лингвистических переменных, термов "
        "или классов и может быть использована в системах поддержки "
        "принятия решений в задачах ресурсосбережения и планирования "
        "эксплуатации автомобильного транспорта.")

    # Источники
    add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    sources = [
        "1. Заде, Л. Понятие лингвистической переменной и его применение "
        "к принятию приближённых решений / Л. Заде. — М. : Мир, 1976. — 165 с.",
        "2. Борисов, А. Н. Принятие решений на основе нечётких моделей / "
        "А. Н. Борисов, О. А. Крумберг, И. П. Фёдоров. — Рига : Зинатне, "
        "1990. — 184 с.",
        "3. Кофман, А. Введение в теорию нечётких множеств / А. Кофман. — "
        "М. : Радио и связь, 1982. — 432 с.",
        "4. Леоненков, А. В. Нечёткое моделирование в среде MATLAB и "
        "fuzzyTECH / А. В. Леоненков. — СПб. : БХВ-Петербург, 2005. — 736 с.",
        "5. Круглов, В. В. Нечёткая логика и искусственные нейронные "
        "сети / В. В. Круглов, М. И. Дли, Р. Ю. Голунов. — М. : Физматлит, "
        "2001. — 224 с.",
        "6. Орлов, А. И. Теория принятия решений / А. И. Орлов. — М. : "
        "Экзамен, 2006. — 573 с.",
        "7. Рутковская, Д. Нейронные сети, генетические алгоритмы и "
        "нечёткие системы / Д. Рутковская, М. Пилиньский, Л. Рутковский. "
        "— М. : Горячая линия — Телеком, 2006. — 452 с.",
    ]
    for s in sources:
        add_par(doc, s, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=Cm(1.25), line_spacing=1.5)

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
